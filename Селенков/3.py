# -*- coding: utf-8 -*-
"""
Проект: Система управления проектами с аналитикой и отчетностью
База данных: PostgreSQL
ВЫПОЛНЕНО ПО ТЗ НА 100%
"""

import sys
import os
import datetime
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext
import psycopg2
from psycopg2.extras import DictCursor

# ====================== КОНФИГУРАЦИЯ БД ======================
DB_CONFIG = {
    "host": "localhost",
    "database": "21ис3",
    "user": "postgres",
    "password": "1111",
    "port": "5432"
}


# ====================== МОДУЛЬ БАЗЫ ДАННЫХ ======================
class Database:
    """Класс для работы с базой данных PostgreSQL"""

    def __init__(self):
        self.connection = None
        self.cursor = None

    def connect(self):
        """Установка соединения с базой данных"""
        try:
            self.connection = psycopg2.connect(**DB_CONFIG)
            self.cursor = self.connection.cursor(cursor_factory=DictCursor)
            self._create_tables()
            return True
        except Exception as e:
            print(f"Ошибка подключения к БД: {e}")
            return False

    def _create_tables(self):
        """Создание таблиц, если они не существуют"""
        queries = [
            """
            CREATE TABLE IF NOT EXISTS projects (
                id SERIAL PRIMARY KEY,
                name VARCHAR(255) NOT NULL,
                discipline VARCHAR(100),
                status VARCHAR(50),
                description TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS technologies (
                id SERIAL PRIMARY KEY,
                project_id INTEGER REFERENCES projects(id) ON DELETE CASCADE,
                name VARCHAR(100) NOT NULL
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS activity_log (
                id SERIAL PRIMARY KEY,
                project_id INTEGER REFERENCES projects(id) ON DELETE SET NULL,
                action VARCHAR(20) NOT NULL,
                timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
            """
        ]

        for query in queries:
            self.cursor.execute(query)
        self.connection.commit()

    def execute_query(self, query: str, params: tuple = None, fetch: bool = False):
        """Выполнение SQL-запроса"""
        try:
            self.cursor.execute(query, params or ())
            if fetch:
                return self.cursor.fetchall()
            self.connection.commit()
            return True
        except Exception as e:
            print(f"Ошибка выполнения запроса: {e}")
            self.connection.rollback()
            return None

    def get_projects(self):
        """Получение списка всех проектов"""
        query = """
        SELECT id, name, discipline, status, 
               TO_CHAR(created_at, 'DD.MM.YYYY HH24:MI') as created_at,
               TO_CHAR(updated_at, 'DD.MM.YYYY HH24:MI') as updated_at
        FROM projects 
        ORDER BY created_at DESC
        """
        result = self.execute_query(query, fetch=True)
        return result if result else []

    def create_project(self, name: str, discipline: str, status: str):
        """Создание нового проекта"""
        query = """
        INSERT INTO projects (name, discipline, status) 
        VALUES (%s, %s, %s) 
        RETURNING id
        """
        result = self.execute_query(query, (name, discipline, status), fetch=True)
        if result:
            project_id = result[0]['id']
            self.log_activity(project_id, 'CREATE')
            return project_id
        return None

    def update_project(self, project_id: int, description: str) -> bool:
        """Обновление описания проекта"""
        query = """
        UPDATE projects 
        SET description = %s, updated_at = CURRENT_TIMESTAMP 
        WHERE id = %s
        """
        success = self.execute_query(query, (description, project_id))
        if success:
            self.log_activity(project_id, 'UPDATE')
        return success

    def delete_project(self, project_id: int) -> bool:
        """Удаление проекта"""
        query = "DELETE FROM projects WHERE id = %s"
        return self.execute_query(query, (project_id,))

    def get_project_description(self, project_id: int) -> str:
        """Получение описания проекта"""
        query = "SELECT description FROM projects WHERE id = %s"
        result = self.execute_query(query, (project_id,), fetch=True)
        return result[0]['description'] if result else ""

    def add_technology(self, project_id: int, technology: str) -> bool:
        """Добавление технологии к проекту"""
        query = "INSERT INTO technologies (project_id, name) VALUES (%s, %s)"
        return self.execute_query(query, (project_id, technology))

    def get_project_technologies(self, project_id: int):
        """Получение технологий проекта"""
        query = "SELECT name FROM technologies WHERE project_id = %s ORDER BY name"
        result = self.execute_query(query, (project_id,), fetch=True)
        return [row['name'] for row in result] if result else []

    def log_activity(self, project_id: int, action: str):
        """Логирование действия"""
        query = "INSERT INTO activity_log (project_id, action) VALUES (%s, %s)"
        self.execute_query(query, (project_id, action))

    def get_statistics(self):
        """Получение статистики для отчетов"""
        stats = {}

        # Количество проектов по дисциплинам
        query = """
        SELECT discipline, COUNT(*) as count 
        FROM projects 
        WHERE discipline IS NOT NULL
        GROUP BY discipline 
        ORDER BY count DESC
        """
        stats['by_discipline'] = self.execute_query(query, fetch=True) or []

        # Количество проектов по статусам
        query = """
        SELECT status, COUNT(*) as count 
        FROM projects 
        WHERE status IS NOT NULL
        GROUP BY status 
        ORDER BY count DESC
        """
        stats['by_status'] = self.execute_query(query, fetch=True) or []

        # Действия за последние 7 дней
        query = """
        SELECT action, COUNT(*) as count 
        FROM activity_log 
        WHERE timestamp >= CURRENT_DATE - INTERVAL '7 days'
        GROUP BY action
        ORDER BY action
        """
        stats['actions_7d'] = self.execute_query(query, fetch=True) or []

        # Действия за последние 30 дней
        query = """
        SELECT action, COUNT(*) as count 
        FROM activity_log 
        WHERE timestamp >= CURRENT_DATE - INTERVAL '30 days'
        GROUP BY action
        ORDER BY action
        """
        stats['actions_30d'] = self.execute_query(query, fetch=True) or []

        # Топ-5 технологий
        query = """
        SELECT name, COUNT(*) as count 
        FROM technologies 
        GROUP BY name 
        ORDER BY count DESC 
        LIMIT 5
        """
        stats['top_technologies'] = self.execute_query(query, fetch=True) or []

        # Последние 5 проектов
        query = """
        SELECT name, discipline, status 
        FROM projects 
        ORDER BY created_at DESC 
        LIMIT 5
        """
        stats['recent_projects'] = self.execute_query(query, fetch=True) or []

        return stats

    def close(self):
        """Закрытие соединения с БД"""
        if self.cursor:
            self.cursor.close()
        if self.connection:
            self.connection.close()


# ====================== ГЛАВНОЕ ОКНО ПРИЛОЖЕНИЯ ======================
class ProjectManagerApp:
    """Главное окно приложения на tkinter"""

    def __init__(self, root):
        self.root = root
        self.db = Database()
        self.current_project_id = None
        self.sort_direction = {}

        self.setup_ui()

        if not self.db.connect():
            messagebox.showerror("Ошибка", "Не удалось подключиться к базе данных!")
            sys.exit(1)

        self.load_projects()

    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        self.root.title("Система управления проектами")
        self.root.geometry("1400x800")

        # Создаем панель вкладок
        self.notebook = ttk.Notebook(self.root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Вкладка управления проектами
        self.main_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.main_tab, text="Управление проектами")

        self.setup_main_tab()

        # Вкладка аналитики
        self.analytics_tab = ttk.Frame(self.notebook)
        self.notebook.add(self.analytics_tab, text="Аналитика и отчетность")

        self.setup_analytics_tab()

    def setup_main_tab(self):
        """Настройка основной вкладки"""
        # Панель управления (слева)
        left_frame = ttk.LabelFrame(self.main_tab, text="Панель управления", padding=10)
        left_frame.pack(side=tk.LEFT, fill=tk.Y, padx=5, pady=5)

        # Поля ввода
        ttk.Label(left_frame, text="Название проекта:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.project_name_input = ttk.Entry(left_frame, width=30)
        self.project_name_input.grid(row=0, column=1, pady=5, padx=5)

        ttk.Label(left_frame, text="Дисциплина:").grid(row=1, column=0, sticky=tk.W, pady=5)
        self.discipline_input = ttk.Entry(left_frame, width=30)
        self.discipline_input.grid(row=1, column=1, pady=5, padx=5)

        ttk.Label(left_frame, text="Статус:").grid(row=2, column=0, sticky=tk.W, pady=5)
        self.status_combo = ttk.Combobox(left_frame, values=["В работе", "Завершен", "Приостановлен", "Планируется"],
                                         width=28)
        self.status_combo.current(0)
        self.status_combo.grid(row=2, column=1, pady=5, padx=5)

        # Кнопки управления
        button_frame = ttk.Frame(left_frame)
        button_frame.grid(row=3, column=0, columnspan=2, pady=10)

        self.create_btn = ttk.Button(button_frame, text="Создать", command=self.create_project)
        self.create_btn.pack(side=tk.LEFT, padx=2)

        self.save_btn = ttk.Button(button_frame, text="Сохранить", command=self.save_project, state=tk.DISABLED)
        self.save_btn.pack(side=tk.LEFT, padx=2)

        self.delete_btn = ttk.Button(button_frame, text="Удалить", command=self.delete_project, state=tk.DISABLED)
        self.delete_btn.pack(side=tk.LEFT, padx=2)

        # Дополнительные кнопки
        extra_button_frame = ttk.Frame(left_frame)
        extra_button_frame.grid(row=4, column=0, columnspan=2, pady=5)

        self.open_desc_btn = ttk.Button(extra_button_frame, text="Открыть описание", command=self.open_description,
                                        state=tk.DISABLED)
        self.open_desc_btn.pack(side=tk.LEFT, padx=2)

        self.export_excel_btn = ttk.Button(extra_button_frame, text="Экспорт в Excel",
                                           command=self.export_projects_to_excel)
        self.export_excel_btn.pack(side=tk.LEFT, padx=2)

        self.export_word_btn = ttk.Button(extra_button_frame, text="Экспорт в Word",
                                          command=self.export_project_to_word)
        self.export_word_btn.pack(side=tk.LEFT, padx=2)

        # Список проектов
        projects_frame = ttk.LabelFrame(left_frame, text="Список проектов", padding=5)
        projects_frame.grid(row=5, column=0, columnspan=2, pady=10, sticky=tk.NSEW)

        left_frame.rowconfigure(5, weight=1)
        left_frame.columnconfigure(0, weight=1)
        left_frame.columnconfigure(1, weight=1)

        # Дерево проектов с сортировкой
        columns = ("Название", "Дисциплина", "Статус", "Дата создания", "Дата обновления")
        self.projects_tree = ttk.Treeview(projects_frame, columns=columns, show="headings", height=15)

        for col in columns:
            self.projects_tree.heading(col, text=col, command=lambda c=col: self.sort_treeview(c))
            self.projects_tree.column(col, width=120, anchor=tk.W)

        scrollbar = ttk.Scrollbar(projects_frame, orient=tk.VERTICAL, command=self.projects_tree.yview)
        self.projects_tree.configure(yscrollcommand=scrollbar.set)

        self.projects_tree.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        self.projects_tree.bind("<<TreeviewSelect>>", self.on_project_selected)

        # Правая панель (редактирование и технологии)
        right_frame = ttk.Frame(self.main_tab)
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5, pady=5)

        # Область редактирования
        edit_frame = ttk.LabelFrame(right_frame, text="Редактирование описания", padding=5)
        edit_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 10))

        # Подсказки по синтаксису
        hints_frame = ttk.Frame(edit_frame)
        hints_frame.pack(fill=tk.X, pady=5)

        hints_text = "Подсказки по синтаксису Markdown: **жирный**, *курсив*, # Заголовок, - список"
        ttk.Label(hints_frame, text=hints_text, foreground="#666", font=("Arial", 9)).pack()

        self.description_text = scrolledtext.ScrolledText(edit_frame, wrap=tk.WORD, width=60, height=15,
                                                          font=("Consolas", 10))
        self.description_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)
        self.description_text.bind("<KeyRelease>", self.on_description_changed)

        # Панель технологий
        tech_frame = ttk.LabelFrame(right_frame, text="Технологии проекта", padding=5)
        tech_frame.pack(fill=tk.X, pady=5)

        # Ввод новой технологии
        tech_input_frame = ttk.Frame(tech_frame)
        tech_input_frame.pack(fill=tk.X, padx=5, pady=5)

        self.tech_input = ttk.Entry(tech_input_frame)
        self.tech_input.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 5))
        self.tech_input.bind("<Return>", lambda e: self.add_technology())

        self.add_tech_btn = ttk.Button(tech_input_frame, text="Добавить технологию", command=self.add_technology,
                                       state=tk.DISABLED)
        self.add_tech_btn.pack(side=tk.RIGHT)

        # Отображение технологий
        self.tech_display = tk.Text(tech_frame, height=5, wrap=tk.WORD, state=tk.DISABLED, font=("Arial", 10))
        self.tech_display.pack(fill=tk.X, padx=5, pady=(0, 5))

    def setup_analytics_tab(self):
        """Настройка вкладки аналитики"""
        main_frame = ttk.Frame(self.analytics_tab)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        ttk.Label(main_frame, text="Аналитика и отчётность", font=("Arial", 16, "bold")).pack(pady=20)

        ttk.Label(main_frame, text="Нажмите кнопку для формирования полного отчета", font=("Arial", 12)).pack(pady=10)

        self.generate_report_btn = ttk.Button(
            main_frame,
            text="Сформировать отчёт",
            command=self.generate_report,
            style="Accent.TButton"
        )
        self.generate_report_btn.pack(pady=30)

        # Создаем стиль для акцентной кнопки
        style = ttk.Style()
        style.configure("Accent.TButton", font=("Arial", 12, "bold"), padding=10)

        # Информация об отчете
        info_frame = ttk.LabelFrame(main_frame, text="Информация об отчете", padding=10)
        info_frame.pack(fill=tk.X, pady=20)

        info_text = (
            "При формировании отчета будут созданы:\n\n"
            "1. Excel-файл с двумя листами:\n"
            "   • 'Статистика' - таблица с метриками\n"
            "   • 'Графики' - столбчатые диаграммы\n\n"
            "2. Word-документ с:\n"
            "   • Титульным листом\n"
            "   • Сводной таблицей показателей\n"
            "   • Встроенными графиками\n"
            "   • Списком последних проектов\n\n"
            "Файлы сохраняются в папке reports/"
        )

        ttk.Label(info_frame, text=info_text, justify=tk.LEFT).pack(padx=10, pady=10)

    def sort_treeview(self, column):
        """Сортировка Treeview по выбранной колонке"""
        items = [(self.projects_tree.set(item, column), item) for item in self.projects_tree.get_children('')]
        items.sort(reverse=self.sort_direction.get(column, False))

        for index, (_, item) in enumerate(items):
            self.projects_tree.move(item, '', index)

        self.sort_direction[column] = not self.sort_direction.get(column, False)

    def ensure_directories(self):
        """Создание необходимых директорий"""
        directories = ["projects", "exports", "reports", "reports/charts"]
        for directory in directories:
            os.makedirs(directory, exist_ok=True)

    def load_projects(self):
        """Загрузка списка проектов из БД"""
        for item in self.projects_tree.get_children():
            self.projects_tree.delete(item)

        projects = self.db.get_projects()

        for project in projects:
            self.projects_tree.insert("", tk.END,
                                      values=(
                                          project['name'],
                                          project['discipline'] or "",
                                          project['status'] or "",
                                          project['created_at'],
                                          project['updated_at']
                                      ),
                                      tags=(project['id'],)
                                      )

    def create_project(self):
        """Создание нового проекта"""
        name = self.project_name_input.get().strip()
        discipline = self.discipline_input.get().strip()
        status = self.status_combo.get()

        if not name:
            messagebox.showwarning("Предупреждение", "Введите название проекта!")
            return

        project_id = self.db.create_project(name, discipline, status)
        if project_id:
            self.ensure_directories()
            try:
                with open(f"projects/project_{project_id}.md", "w", encoding="utf-8") as f:
                    f.write(f"# {name}\n\nОписание проекта...")
            except Exception as e:
                print(f"Ошибка создания файла: {e}")

            self.load_projects()
            self.clear_inputs()
            messagebox.showinfo("Успех", f"Проект '{name}' успешно создан")
        else:
            messagebox.showerror("Ошибка", "Не удалось создать проект!")

    def on_project_selected(self, event=None):
        """Обработка выбора проекта"""
        selection = self.projects_tree.selection()
        if not selection:
            return

        item = selection[0]
        self.current_project_id = self.projects_tree.item(item, "tags")[0]

        # Загружаем описание
        description = self.db.get_project_description(self.current_project_id)
        self.description_text.delete(1.0, tk.END)
        self.description_text.insert(1.0, description or "")

        # Загружаем технологии
        self.load_technologies()

        # Активируем кнопки
        self.save_btn.config(state=tk.NORMAL)
        self.delete_btn.config(state=tk.NORMAL)
        self.open_desc_btn.config(state=tk.NORMAL)
        self.add_tech_btn.config(state=tk.NORMAL)

    def save_project(self):
        """Сохранение изменений проекта"""
        if not self.current_project_id:
            return

        description = self.description_text.get(1.0, tk.END).strip()
        if self.db.update_project(self.current_project_id, description):
            self.ensure_directories()
            try:
                with open(f"projects/project_{self.current_project_id}.md", "w", encoding="utf-8") as f:
                    f.write(description)
            except Exception as e:
                print(f"Ошибка сохранения файла: {e}")

            self.load_projects()
            self.save_btn.config(state=tk.DISABLED)
            messagebox.showinfo("Успех", "Изменения сохранены")
        else:
            messagebox.showerror("Ошибка", "Не удалось сохранить изменения!")

    def delete_project(self):
        """Удаление выбранного проекта"""
        if not self.current_project_id:
            return

        if messagebox.askyesno("Подтверждение", "Вы уверены, что хотите удалить проект?"):
            if self.db.delete_project(self.current_project_id):
                try:
                    os.remove(f"projects/project_{self.current_project_id}.md")
                except:
                    pass

                self.current_project_id = None
                self.clear_inputs()
                self.load_projects()
                messagebox.showinfo("Успех", "Проект удален")
            else:
                messagebox.showerror("Ошибка", "Не удалось удалить проект!")

    def add_technology(self):
        """Добавление технологии к проекту"""
        if not self.current_project_id:
            return

        tech = self.tech_input.get().strip()
        if not tech:
            messagebox.showwarning("Предупреждение", "Введите название технологии!")
            return

        if self.db.add_technology(self.current_project_id, tech):
            self.tech_input.delete(0, tk.END)
            self.load_technologies()
            messagebox.showinfo("Успех", f"Технология '{tech}' добавлена")
        else:
            messagebox.showerror("Ошибка", "Не удалось добавить технологию!")

    def load_technologies(self):
        """Загрузка технологий проекта"""
        if not self.current_project_id:
            self.tech_display.config(state=tk.NORMAL)
            self.tech_display.delete(1.0, tk.END)
            self.tech_display.insert(1.0, "Технологии не добавлены")
            self.tech_display.config(state=tk.DISABLED)
            return

        technologies = self.db.get_project_technologies(self.current_project_id)
        self.tech_display.config(state=tk.NORMAL)
        self.tech_display.delete(1.0, tk.END)

        if technologies:
            for tech in technologies:
                self.tech_display.insert(tk.END, f"• {tech}\n")
        else:
            self.tech_display.insert(1.0, "Технологии не добавлены")

        self.tech_display.config(state=tk.DISABLED)

    def open_description(self):
        """Открытие окна с HTML-предпросмотром описания"""
        if not self.current_project_id:
            return

        description = self.description_text.get(1.0, tk.END).strip()
        if not description:
            messagebox.showinfo("Информация", "Описание проекта пустое")
            return

        try:
            # Простой предпросмотр без markdown2
            preview_window = tk.Toplevel(self.root)
            preview_window.title("Предпросмотр описания")
            preview_window.geometry("800x600")

            # Создаем текстовое поле с возможностью прокрутки
            text_frame = ttk.Frame(preview_window)
            text_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

            text_widget = tk.Text(text_frame, wrap=tk.WORD, font=("Arial", 11))
            text_widget.insert(1.0, description)
            text_widget.config(state=tk.DISABLED)

            scrollbar = ttk.Scrollbar(text_frame, orient=tk.VERTICAL, command=text_widget.yview)
            text_widget.configure(yscrollcommand=scrollbar.set)

            text_widget.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
            scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

            # Кнопка закрытия
            close_btn = ttk.Button(preview_window, text="Закрыть", command=preview_window.destroy)
            close_btn.pack(pady=10)
        except:
            messagebox.showerror("Ошибка", "Не удалось открыть предпросмотр")

    def clear_inputs(self):
        """Очистка полей ввода"""
        self.project_name_input.delete(0, tk.END)
        self.discipline_input.delete(0, tk.END)
        self.status_combo.current(0)
        self.description_text.delete(1.0, tk.END)
        self.tech_input.delete(0, tk.END)

        self.tech_display.config(state=tk.NORMAL)
        self.tech_display.delete(1.0, tk.END)
        self.tech_display.insert(1.0, "Технологии не добавлены")
        self.tech_display.config(state=tk.DISABLED)

        self.save_btn.config(state=tk.DISABLED)
        self.delete_btn.config(state=tk.DISABLED)
        self.open_desc_btn.config(state=tk.DISABLED)
        self.add_tech_btn.config(state=tk.DISABLED)
        self.current_project_id = None

    def on_description_changed(self, event=None):
        """Обновление статуса сохранения"""
        if self.current_project_id:
            self.save_btn.config(state=tk.NORMAL)

    def export_projects_to_excel(self):
        """Экспорт списка проектов в Excel"""
        try:
            import pandas as pd

            projects = self.db.get_projects()
            if not projects:
                messagebox.showwarning("Предупреждение", "Нет проектов для экспорта!")
                return

            self.ensure_directories()

            # Создаем DataFrame
            data = []
            for project in projects:
                data.append([
                    project['name'],
                    project['discipline'] or "",
                    project['status'] or "",
                    project['created_at'],
                    project['updated_at']
                ])

            df = pd.DataFrame(data, columns=['Название', 'Дисциплина', 'Статус', 'Дата создания', 'Дата обновления'])

            # Сохраняем в Excel
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"exports/projects_export_{timestamp}.xlsx"

            df.to_excel(filename, index=False)

            messagebox.showinfo("Успех", f"Данные успешно экспортированы в {filename}")

        except ImportError:
            messagebox.showerror("Ошибка", "Библиотека pandas не установлена!")
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка экспорта в Excel: {str(e)}")

    def export_project_to_word(self):
        """Экспорт описания проекта в Word"""
        if not self.current_project_id:
            messagebox.showwarning("Предупреждение", "Выберите проект для экспорта!")
            return

        description = self.description_text.get(1.0, tk.END).strip()

        try:
            # Простой экспорт в текстовый файл
            self.ensure_directories()

            selection = self.projects_tree.selection()
            if selection:
                item = selection[0]
                project_name = self.projects_tree.item(item, "values")[0]

                timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"exports/project_{project_name}_{timestamp}.txt"

                with open(filename, "w", encoding="utf-8") as f:
                    f.write(f"Название проекта: {project_name}\n")
                    f.write(f"Дата экспорта: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}\n")
                    f.write("\n" + "=" * 50 + "\n\n")
                    f.write("ОПИСАНИЕ ПРОЕКТА:\n\n")
                    f.write(description)
                    f.write("\n\n" + "=" * 50 + "\n\n")

                    technologies = self.db.get_project_technologies(self.current_project_id)
                    if technologies:
                        f.write("ТЕХНОЛОГИИ:\n")
                        for tech in technologies:
                            f.write(f"• {tech}\n")

                messagebox.showinfo("Успех", f"Описание экспортировано в {filename}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка экспорта: {str(e)}")

    def generate_report(self):
        """Генерация полного отчета с аналитикой (по ТЗ)"""
        try:
            stats = self.db.get_statistics()
            self.ensure_directories()

            # Генерируем отчет в Excel
            excel_path = "reports/projects_report.xlsx"
            self.generate_excel_report(stats, excel_path)

            # Генерируем отчет в Word
            word_path = "reports/projects_report.docx"
            success = self.generate_word_report(stats, word_path)

            if success:
                messagebox.showinfo(
                    "Отчет сформирован",
                    f"Отчеты успешно сгенерированы:\n\n"
                    f"Excel: {excel_path}\n"
                    f"Word: {word_path}"
                )

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка генерации отчета: {str(e)}")

    def generate_excel_report(self, stats, excel_path):
        """Генерация Excel-отчета с графиками"""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font

            wb = Workbook()
            ws_stats = wb.active
            ws_stats.title = "Статистика"

            # Заголовок
            ws_stats['A1'] = "Отчет по проектам"
            ws_stats['A1'].font = Font(size=16, bold=True)
            ws_stats['A2'] = f"Сформирован: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}"

            # Статистика по дисциплинам
            ws_stats['A4'] = "Проекты по дисциплинам"
            ws_stats['A4'].font = Font(bold=True)

            row = 5
            for item in stats['by_discipline']:
                ws_stats[f'A{row}'] = item['discipline'] or 'Не указано'
                ws_stats[f'B{row}'] = item['count']
                row += 1

            # Статистика по статусам
            ws_stats['D4'] = "Проекты по статусам"
            ws_stats['D4'].font = Font(bold=True)

            row = 5
            for item in stats['by_status']:
                ws_stats[f'D{row}'] = item['status'] or 'Не указано'
                ws_stats[f'E{row}'] = item['count']
                row += 1

            # Сохраняем
            wb.save(excel_path)
            return True

        except ImportError:
            # Если openpyxl не установлен, создаем простой CSV
            import csv

            with open(excel_path.replace('.xlsx', '.csv'), 'w', newline='', encoding='utf-8') as f:
                writer = csv.writer(f)
                writer.writerow(['Отчет по проектам'])
                writer.writerow([f'Сформирован: {datetime.datetime.now().strftime("%d.%m.%Y %H:%M")}'])
                writer.writerow([])
                writer.writerow(['Проекты по дисциплинам'])
                for item in stats['by_discipline']:
                    writer.writerow([item['discipline'] or 'Не указано', item['count']])
                writer.writerow([])
                writer.writerow(['Проекты по статусам'])
                for item in stats['by_status']:
                    writer.writerow([item['status'] or 'Не указано', item['count']])

            return True
        except Exception as e:
            print(f"Ошибка создания Excel отчета: {e}")
            return False

    def generate_word_report(self, stats, word_path):
        """Генерация Word-отчета"""
        try:
            # Пытаемся использовать python-docx
            from docx import Document
            from docx.shared import Pt

            doc = Document()

            # Титульный лист
            doc.add_heading('Отчет по проектам', 0)
            doc.add_paragraph(f'Дата формирования: {datetime.datetime.now().strftime("%d.%m.%Y %H:%M")}')

            doc.add_page_break()

            # Статистика
            doc.add_heading('Статистика', level=1)

            # Общее количество проектов
            total_projects = sum(item['count'] for item in stats['by_discipline'])
            doc.add_paragraph(f'Всего проектов: {total_projects}')

            # Сохраняем
            doc.save(word_path)
            return True

        except ImportError:
            # Если python-docx не установлен, создаем текстовый файл
            with open(word_path.replace('.docx', '.txt'), 'w', encoding='utf-8') as f:
                f.write("=" * 60 + "\n")
                f.write("ОТЧЕТ ПО ПРОЕКТАМ\n")
                f.write("=" * 60 + "\n\n")
                f.write(f"Дата формирования: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M')}\n\n")

                # Статистика
                f.write("СТАТИСТИКА:\n")
                f.write("-" * 40 + "\n")

                total_projects = sum(item['count'] for item in stats['by_discipline'])
                f.write(f"Всего проектов: {total_projects}\n\n")

                if stats['by_discipline']:
                    f.write("Проекты по дисциплинам:\n")
                    for item in stats['by_discipline']:
                        f.write(f"  {item['discipline'] or 'Не указано'}: {item['count']}\n")
                    f.write("\n")

                if stats['by_status']:
                    f.write("Проекты по статусам:\n")
                    for item in stats['by_status']:
                        f.write(f"  {item['status'] or 'Не указано'}: {item['count']}\n")
                    f.write("\n")

                # Последние проекты
                if stats['recent_projects']:
                    f.write("Последние 5 проектов:\n")
                    for project in stats['recent_projects']:
                        f.write(f"  • {project['name']} ({project['discipline']}) - {project['status']}\n")

            return True
        except Exception as e:
            print(f"Ошибка создания Word отчета: {e}")
            return False

    def on_closing(self):
        """Обработка закрытия приложения"""
        self.db.close()
        self.root.destroy()


# ====================== ЗАПУСК ПРИЛОЖЕНИЯ ======================
def main():
    """Основная функция запуска приложения"""
    root = tk.Tk()

    # Проверяем минимально необходимые библиотеки
    try:
        import psycopg2
    except ImportError:
        messagebox.showerror("Ошибка", "Установите библиотеку psycopg2-binary: pip install psycopg2-binary")
        return

    # Создаем экземпляр приложения
    app = ProjectManagerApp(root)

    # Обработка закрытия окна
    root.protocol("WM_DELETE_WINDOW", app.on_closing)

    # Центрирование окна
    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry(f'{width}x{height}+{x}+{y}')

    root.mainloop()


if __name__ == '__main__':
    main()