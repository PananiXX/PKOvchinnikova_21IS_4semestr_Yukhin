# portfolio_manager.py
import os
import sys
import tkinter as tk
from tkinter import ttk, messagebox, scrolledtext, filedialog
import datetime
from pathlib import Path
import markdown
import webbrowser
import json
import traceback
from typing import Optional, List, Dict, Any, Tuple
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure
import pandas as pd
from openpyxl import Workbook
from openpyxl.drawing.image import Image as ExcelImage
from openpyxl.chart import BarChart, Reference
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
import matplotlib

matplotlib.use('Agg')

# Попытка импорта PostgreSQL, если не доступен - используем SQLite
try:
    import psycopg2
    from psycopg2.extras import RealDictCursor

    POSTGRES_AVAILABLE = True
except ImportError:
    POSTGRES_AVAILABLE = False
    import sqlite3

    print("PostgreSQL не установлен, используется SQLite")


class DatabaseManager:
    """Универсальный менеджер базы данных"""

    def __init__(self):
        self.connection = None
        self.cursor = None
        self.db_type = None
        self.connect()

    def connect(self):
        """Подключение к базе данных"""
        try:
            if POSTGRES_AVAILABLE:
                # Пробуем подключиться к PostgreSQL
                self.connection = psycopg2.connect(
                    host="localhost",
                    database="21ис5",
                    user="postgres",
                    password="1111",
                    port="5432"
                )
                self.cursor = self.connection.cursor()
                self.db_type = 'postgresql'
                print("Успешное подключение к PostgreSQL")
                self.create_tables()
            else:
                raise ConnectionError("PostgreSQL не доступен")

        except Exception as e:
            print(f"PostgreSQL недоступен: {e}, используется SQLite")
            # Используем SQLite как резервную базу
            self.connection = sqlite3.connect('portfolio.db', check_same_thread=False)
            self.connection.row_factory = sqlite3.Row
            self.cursor = self.connection.cursor()
            self.db_type = 'sqlite'
            self.create_tables()

    def create_tables(self):
        """Создание таблиц"""
        try:
            if self.db_type == 'postgresql':
                # PostgreSQL схемы
                self.cursor.execute('''
                    CREATE TABLE IF NOT EXISTS records (
                        id SERIAL PRIMARY KEY,
                        title VARCHAR(255) NOT NULL,
                        type VARCHAR(100) NOT NULL,
                        year INTEGER NOT NULL,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        description TEXT,
                        file_path VARCHAR(500)
                    )
                ''')

                self.cursor.execute('''
                    CREATE TABLE IF NOT EXISTS coauthors (
                        id SERIAL PRIMARY KEY,
                        record_id INTEGER REFERENCES records(id) ON DELETE CASCADE,
                        name VARCHAR(255) NOT NULL
                    )
                ''')

                self.cursor.execute('''
                    CREATE TABLE IF NOT EXISTS activity_log (
                        id SERIAL PRIMARY KEY,
                        action VARCHAR(100) NOT NULL,
                        record_id INTEGER,
                        timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        details TEXT
                    )
                ''')

            else:  # SQLite
                self.cursor.execute('''
                    CREATE TABLE IF NOT EXISTS records (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        title TEXT NOT NULL,
                        type TEXT NOT NULL,
                        year INTEGER NOT NULL,
                        created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        description TEXT,
                        file_path TEXT
                    )
                ''')

                self.cursor.execute('''
                    CREATE TABLE IF NOT EXISTS coauthors (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        record_id INTEGER REFERENCES records(id) ON DELETE CASCADE,
                        name TEXT NOT NULL
                    )
                ''')

                self.cursor.execute('''
                    CREATE TABLE IF NOT EXISTS activity_log (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        action TEXT NOT NULL,
                        record_id INTEGER,
                        timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                        details TEXT
                    )
                ''')

            self.connection.commit()
            print("Таблицы успешно созданы")

        except Exception as e:
            print(f"Ошибка создания таблиц: {e}")
            traceback.print_exc()

    def execute_query(self, query: str, params: tuple = (), fetch: bool = False, fetch_one: bool = False):
        """Выполнение SQL запроса"""
        try:
            if self.db_type == 'postgresql':
                self.cursor.execute(query, params)
            else:
                # Адаптация параметров для SQLite
                query = query.replace('%s', '?')
                self.cursor.execute(query, params)

            if fetch:
                if self.db_type == 'postgresql':
                    return [dict(row) for row in self.cursor.fetchall()]
                else:
                    return [dict(row) for row in self.cursor.fetchall()]
            elif fetch_one:
                if self.db_type == 'postgresql':
                    result = self.cursor.fetchone()
                    return dict(result) if result else None
                else:
                    result = self.cursor.fetchone()
                    return dict(result) if result else None
            else:
                self.connection.commit()
                return True

        except Exception as e:
            print(f"Ошибка выполнения запроса: {e}")
            print(f"Query: {query}")
            print(f"Params: {params}")
            traceback.print_exc()
            return False

    def add_record(self, title: str, record_type: str, year: int, description: str = "") -> Optional[int]:
        """Добавление новой записи"""
        try:
            # Создаем директорию для файлов
            os.makedirs('records', exist_ok=True)

            # Генерируем имя файла
            safe_title = "".join(c if c.isalnum() else "_" for c in title)
            file_name = f"{safe_title}_{year}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
            file_path = os.path.join('records', file_name)

            # Сохраняем описание в файл
            with open(file_path, 'w', encoding='utf-8') as f:
                f.write(description)

            # Вставляем запись в базу данных
            if self.db_type == 'postgresql':
                query = '''
                    INSERT INTO records (title, type, year, description, file_path, created_at, updated_at)
                    VALUES (%s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP)
                    RETURNING id
                '''
            else:
                query = '''
                    INSERT INTO records (title, type, year, description, file_path, created_at, updated_at)
                    VALUES (?, ?, ?, ?, ?, datetime('now'), datetime('now'))
                '''

            params = (title, record_type, year, description, file_path)

            if self.db_type == 'postgresql':
                self.cursor.execute(query, params)
                record_id = self.cursor.fetchone()[0]
            else:
                self.cursor.execute(query, params)
                record_id = self.cursor.lastrowid

            # Логируем действие
            self.execute_query('''
                INSERT INTO activity_log (action, record_id, details)
                VALUES (%s, %s, %s)
            ''', ('CREATE', record_id, f'Создана запись: {title}'))

            self.connection.commit()
            return record_id

        except Exception as e:
            print(f"Ошибка добавления записи: {e}")
            traceback.print_exc()
            return None

    def update_record(self, record_id: int, **kwargs) -> bool:
        """Обновление записи"""
        try:
            if not kwargs:
                return True

            # Получаем текущие данные
            record = self.get_record_by_id(record_id)
            if not record:
                return False

            # Обновляем файл, если изменилось описание
            if 'description' in kwargs and record['file_path']:
                try:
                    with open(record['file_path'], 'w', encoding='utf-8') as f:
                        f.write(kwargs['description'])
                except:
                    pass  # Файл может не существовать

            # Формируем SQL запрос
            set_clause = []
            params = []

            for key, value in kwargs.items():
                if key == 'description' and value is not None:
                    set_clause.append("description = %s")
                    params.append(value)
                elif key == 'title' and value:
                    set_clause.append("title = %s")
                    params.append(value)
                elif key == 'type' and value:
                    set_clause.append("type = %s")
                    params.append(value)
                elif key == 'year' and value:
                    set_clause.append("year = %s")
                    params.append(value)

            if not set_clause:
                return True

            set_clause.append("updated_at = CURRENT_TIMESTAMP")

            if self.db_type == 'postgresql':
                query = f"UPDATE records SET {', '.join(set_clause)} WHERE id = %s"
            else:
                query = f"UPDATE records SET {', '.join(set_clause.replace('%s', '?'))} WHERE id = ?"

            params.append(record_id)

            success = self.execute_query(query, tuple(params))

            if success:
                # Логируем действие
                self.execute_query('''
                    INSERT INTO activity_log (action, record_id, details)
                    VALUES (%s, %s, %s)
                ''', ('UPDATE', record_id, f'Обновлена запись ID: {record_id}'))

            return success

        except Exception as e:
            print(f"Ошибка обновления записи: {e}")
            return False

    def delete_record(self, record_id: int) -> bool:
        """Удаление записи"""
        try:
            # Получаем путь к файлу
            record = self.get_record_by_id(record_id)
            if record and record.get('file_path'):
                try:
                    if os.path.exists(record['file_path']):
                        os.remove(record['file_path'])
                except:
                    pass  # Игнорируем ошибки удаления файла

            # Удаляем запись из базы
            query = "DELETE FROM records WHERE id = %s"
            success = self.execute_query(query, (record_id,))

            if success:
                # Логируем действие
                self.execute_query('''
                    INSERT INTO activity_log (action, record_id, details)
                    VALUES (%s, %s, %s)
                ''', ('DELETE', record_id, f'Удалена запись ID: {record_id}'))

            return success

        except Exception as e:
            print(f"Ошибка удаления записи: {e}")
            return False

    def get_all_records(self) -> List[Dict]:
        """Получение всех записей"""
        query = '''
            SELECT id, title, type, year, created_at, updated_at, description, file_path
            FROM records
            ORDER BY created_at DESC
        '''
        return self.execute_query(query, fetch=True) or []

    def get_record_by_id(self, record_id: int) -> Optional[Dict]:
        """Получение записи по ID"""
        query = '''
            SELECT id, title, type, year, created_at, updated_at, description, file_path
            FROM records
            WHERE id = %s
        '''
        return self.execute_query(query, (record_id,), fetch_one=True)

    def add_coauthor(self, record_id: int, name: str) -> bool:
        """Добавление соавтора"""
        try:
            success = self.execute_query('''
                INSERT INTO coauthors (record_id, name)
                VALUES (%s, %s)
            ''', (record_id, name))

            if success:
                # Логируем действие
                self.execute_query('''
                    INSERT INTO activity_log (action, record_id, details)
                    VALUES (%s, %s, %s)
                ''', ('ADD_COAUTHOR', record_id, f'Добавлен соавтор: {name}'))

            return success

        except Exception as e:
            print(f"Ошибка добавления соавтора: {e}")
            return False

    def get_coauthors(self, record_id: int) -> List[str]:
        """Получение соавторов записи"""
        query = '''
            SELECT name FROM coauthors
            WHERE record_id = %s
            ORDER BY id
        '''
        result = self.execute_query(query, (record_id,), fetch=True)
        return [row['name'] for row in result] if result else []

    def delete_coauthor(self, record_id: int, name: str) -> bool:
        """Удаление соавтора"""
        try:
            success = self.execute_query('''
                DELETE FROM coauthors
                WHERE record_id = %s AND name = %s
            ''', (record_id, name))

            if success:
                # Логируем действие
                self.execute_query('''
                    INSERT INTO activity_log (action, record_id, details)
                    VALUES (%s, %s, %s)
                ''', ('REMOVE_COAUTHOR', record_id, f'Удален соавтор: {name}'))

            return success

        except Exception as e:
            print(f"Ошибка удаления соавтора: {e}")
            return False

    def get_statistics(self) -> Dict:
        """Получение статистики"""
        try:
            stats = {}

            # Общее количество записей
            result = self.execute_query('SELECT COUNT(*) as count FROM records', fetch_one=True)
            stats['total_records'] = result['count'] if result else 0

            # Распределение по типам
            query = '''
                SELECT type, COUNT(*) as count
                FROM records
                GROUP BY type
                ORDER BY count DESC
            '''
            result = self.execute_query(query, fetch=True)
            stats['type_distribution'] = {row['type']: row['count'] for row in result} if result else {}

            # Распределение по годам
            query = '''
                SELECT year, COUNT(*) as count
                FROM records
                GROUP BY year
                ORDER BY year
            '''
            result = self.execute_query(query, fetch=True)
            stats['year_distribution'] = {row['year']: row['count'] for row in result} if result else {}

            # Количество уникальных соавторов
            query = 'SELECT COUNT(DISTINCT name) as count FROM coauthors'
            result = self.execute_query(query, fetch_one=True)
            stats['unique_coauthors'] = result['count'] if result else 0

            # Активность за последние 12 месяцев
            if self.db_type == 'postgresql':
                query = '''
                    SELECT 
                        TO_CHAR(DATE_TRUNC('month', created_at), 'YYYY-MM') as month,
                        COUNT(*) as count
                    FROM records
                    WHERE created_at >= CURRENT_DATE - INTERVAL '12 months'
                    GROUP BY DATE_TRUNC('month', created_at)
                    ORDER BY month
                '''
            else:
                query = '''
                    SELECT 
                        strftime('%Y-%m', created_at) as month,
                        COUNT(*) as count
                    FROM records
                    WHERE created_at >= datetime('now', '-12 months')
                    GROUP BY strftime('%Y-%m', created_at)
                    ORDER BY month
                '''

            result = self.execute_query(query, fetch=True)
            stats['monthly_activity'] = {row['month']: row['count'] for row in result} if result else {}

            return stats

        except Exception as e:
            print(f"Ошибка получения статистики: {e}")
            return {}


class PortfolioApp:
    """Основной класс приложения"""

    def __init__(self, root):
        self.root = root
        self.root.title("Portfolio Management System")
        self.root.geometry("1200x800")

        # Инициализация базы данных
        self.db = DatabaseManager()

        # Текущая выбранная запись
        self.current_record_id = None

        # Настройка иконки
        try:
            self.root.iconbitmap('icon.ico')
        except:
            pass

        # Создание GUI
        self.setup_ui()

        # Загрузка записей
        self.load_records()

        # Создаем необходимые директории
        self.create_directories()

    def create_directories(self):
        """Создание необходимых директорий"""
        os.makedirs('records', exist_ok=True)
        os.makedirs('reports', exist_ok=True)
        os.makedirs('temp', exist_ok=True)

    def setup_ui(self):
        """Настройка пользовательского интерфейса"""
        # Создаем стили
        self.setup_styles()

        # Главный контейнер
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # Конфигурация расширения
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)
        main_frame.rowconfigure(2, weight=1)

        # === Панель управления ===
        control_frame = ttk.LabelFrame(main_frame, text="Панель управления", padding="10")
        control_frame.grid(row=0, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(0, 10))

        # Первая строка: поля ввода
        input_frame = ttk.Frame(control_frame)
        input_frame.pack(fill=tk.X, pady=(0, 10))

        # Название
        ttk.Label(input_frame, text="Название:").grid(row=0, column=0, sticky=tk.W, padx=(0, 5))
        self.title_entry = ttk.Entry(input_frame, width=40)
        self.title_entry.grid(row=0, column=1, padx=(0, 20))

        # Тип записи
        ttk.Label(input_frame, text="Тип записи:").grid(row=0, column=2, sticky=tk.W, padx=(0, 5))
        self.type_combobox = ttk.Combobox(input_frame, width=20, values=[
            "Статья", "Книга", "Доклад", "Патент", "Проект", "Исследование",
            "Курсовая", "Диплом", "Монография", "Отчёт", "Другое"
        ])
        self.type_combobox.grid(row=0, column=3, padx=(0, 20))
        self.type_combobox.set("Статья")

        # Год
        ttk.Label(input_frame, text="Год:").grid(row=0, column=4, sticky=tk.W, padx=(0, 5))
        self.year_spinbox = ttk.Spinbox(input_frame, from_=2000, to=2030, width=10)
        current_year = datetime.datetime.now().year
        self.year_spinbox.set(current_year)
        self.year_spinbox.grid(row=0, column=5)

        # Вторая строка: кнопки
        buttons_frame = ttk.Frame(control_frame)
        buttons_frame.pack(fill=tk.X)

        # Кнопки управления
        buttons = [
            ("Создать", self.create_record, 0),
            ("Сохранить", self.save_record, 1),
            ("Удалить", self.delete_record, 2),
            ("Открыть описание", self.open_description, 3),
            ("Экспорт в Excel", self.export_to_excel, 4),
            ("Экспорт в Word", self.export_to_word, 5),
        ]

        for text, command, col in buttons:
            btn = ttk.Button(buttons_frame, text=text, command=command)
            btn.grid(row=0, column=col, padx=2)

            # Сохраняем ссылки на кнопки, которые нужно отключать
            if text == "Сохранить":
                self.save_btn = btn
                btn.config(state=tk.DISABLED)
            elif text == "Удалить":
                self.delete_btn = btn
                btn.config(state=tk.DISABLED)

        # Кнопка аналитики
        self.analytics_btn = ttk.Button(buttons_frame, text="Аналитика и отчётность",
                                        command=self.open_analytics)
        self.analytics_btn.grid(row=0, column=6, padx=2)

        # === Основное содержимое ===
        content_frame = ttk.Frame(main_frame)
        content_frame.grid(row=1, column=0, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 10))

        # Конфигурация расширения
        content_frame.columnconfigure(1, weight=1)
        content_frame.rowconfigure(0, weight=1)

        # === Левая панель: список записей ===
        left_frame = ttk.LabelFrame(content_frame, text="Список записей", padding="5")
        left_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=(0, 5))

        # Treeview
        columns = ("ID", "Название", "Тип", "Год", "Дата создания")
        self.tree = ttk.Treeview(left_frame, columns=columns, show="headings", height=25)

        # Настройка колонок
        column_widths = {"ID": 50, "Название": 250, "Тип": 100, "Год": 60, "Дата создания": 120}
        for col in columns:
            self.tree.heading(col, text=col,
                              command=lambda c=col: self.sort_treeview(c, False if c == self.last_sort_col else True))
            self.tree.column(col, width=column_widths.get(col, 100))

        self.last_sort_col = None
        self.sort_reverse = False

        # Scrollbar
        tree_scrollbar = ttk.Scrollbar(left_frame, orient=tk.VERTICAL, command=self.tree.yview)
        self.tree.configure(yscrollcommand=tree_scrollbar.set)

        # Размещение
        self.tree.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        tree_scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))

        # Конфигурация расширения
        left_frame.columnconfigure(0, weight=1)
        left_frame.rowconfigure(0, weight=1)

        # Привязка события выбора
        self.tree.bind('<<TreeviewSelect>>', self.on_record_select)

        # === Правая панель: редактирование и соавторы ===
        right_frame = ttk.Frame(content_frame)
        right_frame.grid(row=0, column=1, sticky=(tk.W, tk.E, tk.N, tk.S))

        # Конфигурация расширения
        right_frame.columnconfigure(0, weight=1)
        right_frame.rowconfigure(0, weight=3)
        right_frame.rowconfigure(1, weight=1)

        # Область редактирования
        edit_frame = ttk.LabelFrame(right_frame, text="Редактирование описания (Markdown)", padding="5")
        edit_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), pady=(0, 5))

        # Text виджет для редактирования
        self.text_editor = scrolledtext.ScrolledText(edit_frame, wrap=tk.WORD, width=60, height=20)
        self.text_editor.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # Подсказки по синтаксису
        help_frame = ttk.Frame(edit_frame)
        help_frame.grid(row=1, column=0, sticky=(tk.W, tk.E), pady=(5, 0))

        ttk.Label(help_frame, text="Подсказки: ", font=("Arial", 9)).pack(side=tk.LEFT)
        syntax_hints = [
            "Заголовок: # текст",
            "Цитаты: > текст",
            "Код: ```код```",
            "Ссылки: [текст](url)",
            "Списки: - пункт"
        ]

        for hint in syntax_hints:
            ttk.Label(help_frame, text=hint, font=("Arial", 9),
                      foreground="gray").pack(side=tk.LEFT, padx=5)

        # Конфигурация расширения
        edit_frame.columnconfigure(0, weight=1)
        edit_frame.rowconfigure(0, weight=1)

        # === Панель соавторов ===
        coauthor_frame = ttk.LabelFrame(right_frame, text="Соавторы", padding="10")
        coauthor_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # Верхняя часть: добавление соавторов
        add_frame = ttk.Frame(coauthor_frame)
        add_frame.pack(fill=tk.X, pady=(0, 10))

        ttk.Label(add_frame, text="Имя соавтора:").pack(side=tk.LEFT, padx=(0, 5))
        self.coauthor_entry = ttk.Entry(add_frame, width=30)
        self.coauthor_entry.pack(side=tk.LEFT, padx=(0, 5))

        self.add_coauthor_btn = ttk.Button(add_frame, text="Добавить",
                                           command=self.add_coauthor, state=tk.DISABLED)
        self.add_coauthor_btn.pack(side=tk.LEFT, padx=(0, 5))

        # Кнопка удаления выбранного соавтора
        self.remove_coauthor_btn = ttk.Button(add_frame, text="Удалить выбранного",
                                              command=self.remove_selected_coauthor, state=tk.DISABLED)
        self.remove_coauthor_btn.pack(side=tk.LEFT)

        # Нижняя часть: список соавторов
        list_frame = ttk.Frame(coauthor_frame)
        list_frame.pack(fill=tk.BOTH, expand=True)

        # Listbox для отображения соавторов
        self.coauthors_listbox = tk.Listbox(list_frame, height=5)
        self.coauthors_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Scrollbar для listbox
        list_scrollbar = ttk.Scrollbar(list_frame, orient=tk.VERTICAL,
                                       command=self.coauthors_listbox.yview)
        self.coauthors_listbox.config(yscrollcommand=list_scrollbar.set)
        list_scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        # Привязка события выбора в listbox
        self.coauthors_listbox.bind('<<ListboxSelect>>', self.on_coauthor_select)

        # === Статус бар ===
        self.status_bar = ttk.Label(main_frame, text="Готов к работе", relief=tk.SUNKEN, anchor=tk.W)
        self.status_bar.grid(row=2, column=0, columnspan=2, sticky=(tk.W, tk.E), pady=(5, 0))

    def setup_styles(self):
        """Настройка стилей для виджетов"""
        style = ttk.Style()

        # Настраиваем цвета
        style.configure("Treeview", rowheight=25)
        style.configure("Treeview.Heading", font=('Arial', 10, 'bold'))

        # Акцентная кнопка
        style.configure("Accent.TButton", font=('Arial', 10, 'bold'))

    def load_records(self):
        """Загрузка записей в Treeview"""
        # Очищаем текущий список
        for item in self.tree.get_children():
            self.tree.delete(item)

        # Получаем записи из базы
        records = self.db.get_all_records()

        # Добавляем записи в Treeview
        for record in records:
            created_at = record['created_at']
            if isinstance(created_at, str):
                created_str = created_at[:16]  # Берем первые 16 символов для формата YYYY-MM-DD HH:MM
            elif hasattr(created_at, 'strftime'):
                created_str = created_at.strftime('%Y-%m-%d %H:%M')
            else:
                created_str = str(created_at)

            self.tree.insert('', 'end', values=(
                record['id'],
                record['title'],
                record['type'],
                record['year'],
                created_str
            ))

        self.update_status(f"Загружено записей: {len(records)}")

    def sort_treeview(self, col, reverse):
        """Сортировка Treeview по колонке"""
        # Получаем все элементы
        items = [(self.tree.set(item, col), item) for item in self.tree.get_children('')]

        # Определяем тип сортировки
        if col in ["ID", "Год"]:
            # Числовая сортировка
            try:
                items.sort(key=lambda x: int(x[0]) if x[0] else 0, reverse=reverse)
            except:
                items.sort(key=lambda x: x[0], reverse=reverse)
        elif col == "Дата создания":
            # Сортировка по дате
            try:
                items.sort(key=lambda x: datetime.datetime.strptime(x[0], '%Y-%m-%d %H:%M')
                if x[0] else datetime.datetime.min, reverse=reverse)
            except:
                items.sort(key=lambda x: x[0], reverse=reverse)
        else:
            # Строковая сортировка
            items.sort(key=lambda x: x[0].lower() if x[0] else '', reverse=reverse)

        # Перемещаем элементы
        for index, (_, item) in enumerate(items):
            self.tree.move(item, '', index)

        # Запоминаем параметры сортировки
        self.last_sort_col = col
        self.sort_reverse = reverse

        # Обновляем заголовок
        for column in self.tree['columns']:
            self.tree.heading(column,
                              command=lambda c=column: self.sort_treeview(c, False if c == col else True))

    def on_record_select(self, event):
        """Обработка выбора записи"""
        selection = self.tree.selection()
        if not selection:
            return

        # Получаем ID выбранной записи
        item = self.tree.item(selection[0])
        record_id = item['values'][0]

        # Загружаем данные записи
        record = self.db.get_record_by_id(record_id)
        if not record:
            return

        self.current_record_id = record_id

        # Заполняем поля формы
        self.title_entry.delete(0, tk.END)
        self.title_entry.insert(0, record['title'])

        self.type_combobox.set(record['type'])
        self.year_spinbox.set(record['year'])

        # Загружаем описание
        self.text_editor.delete(1.0, tk.END)
        if record['description']:
            self.text_editor.insert(1.0, record['description'])

        # Загружаем соавторов
        self.load_coauthors(record_id)

        # Активируем кнопки
        self.save_btn.config(state=tk.NORMAL)
        self.delete_btn.config(state=tk.NORMAL)
        self.add_coauthor_btn.config(state=tk.NORMAL)

        self.update_status(f"Выбрана запись: {record['title']}")

    def load_coauthors(self, record_id):
        """Загрузка соавторов записи"""
        # Очищаем список
        self.coauthors_listbox.delete(0, tk.END)

        # Загружаем соавторов
        coauthors = self.db.get_coauthors(record_id)

        # Добавляем в listbox
        for coauthor in coauthors:
            self.coauthors_listbox.insert(tk.END, coauthor)

        # Активируем кнопку удаления если есть соавторы
        if coauthors:
            self.remove_coauthor_btn.config(state=tk.NORMAL)
        else:
            self.remove_coauthor_btn.config(state=tk.DISABLED)

    def on_coauthor_select(self, event):
        """Обработка выбора соавтора"""
        selection = self.coauthors_listbox.curselection()
        if selection:
            self.remove_coauthor_btn.config(state=tk.NORMAL)
        else:
            self.remove_coauthor_btn.config(state=tk.DISABLED)

    def create_record(self):
        """Создание новой записи"""
        # Получаем данные из формы
        title = self.title_entry.get().strip()
        record_type = self.type_combobox.get().strip()
        year_str = self.year_spinbox.get().strip()

        # Валидация
        if not title:
            messagebox.showerror("Ошибка", "Введите название записи")
            self.title_entry.focus()
            return

        if not record_type:
            messagebox.showerror("Ошибка", "Выберите тип записи")
            self.type_combobox.focus()
            return

        try:
            year = int(year_str)
            if year < 2000 or year > 2030:
                raise ValueError
        except ValueError:
            messagebox.showerror("Ошибка", "Введите корректный год (2000-2030)")
            self.year_spinbox.focus()
            return

        # Получаем описание
        description = self.text_editor.get(1.0, tk.END).strip()

        # Создаем запись
        record_id = self.db.add_record(title, record_type, year, description)

        if record_id:
            messagebox.showinfo("Успех", f"Запись успешно создана (ID: {record_id})")

            # Обновляем список записей
            self.load_records()

            # Сбрасываем форму
            self.clear_form()

            self.update_status(f"Создана новая запись: {title}")
        else:
            messagebox.showerror("Ошибка", "Не удалось создать запись")

    def save_record(self):
        """Сохранение изменений записи"""
        if not self.current_record_id:
            return

        # Получаем данные из формы
        title = self.title_entry.get().strip()
        record_type = self.type_combobox.get().strip()
        year_str = self.year_spinbox.get().strip()
        description = self.text_editor.get(1.0, tk.END).strip()

        # Валидация
        if not title:
            messagebox.showerror("Ошибка", "Введите название записи")
            self.title_entry.focus()
            return

        try:
            year = int(year_str)
        except ValueError:
            messagebox.showerror("Ошибка", "Введите корректный год")
            self.year_spinbox.focus()
            return

        # Обновляем запись
        success = self.db.update_record(
            self.current_record_id,
            title=title,
            record_type=record_type,
            year=year,
            description=description
        )

        if success:
            messagebox.showinfo("Успех", "Запись успешно обновлена")
            self.load_records()
            self.update_status(f"Запись обновлена: {title}")
        else:
            messagebox.showerror("Ошибка", "Не удалось обновить запись")

    def delete_record(self):
        """Удаление записи"""
        if not self.current_record_id:
            return

        # Получаем название записи для подтверждения
        record = self.db.get_record_by_id(self.current_record_id)
        if not record:
            return

        # Подтверждение удаления
        if not messagebox.askyesno("Подтверждение",
                                   f"Вы уверены, что хотите удалить запись:\n\"{record['title']}\"?"):
            return

        # Удаляем запись
        success = self.db.delete_record(self.current_record_id)

        if success:
            messagebox.showinfo("Успех", "Запись успешно удалена")
            self.clear_form()
            self.load_records()
            self.update_status("Запись удалена")
        else:
            messagebox.showerror("Ошибка", "Не удалось удалить запись")

    def add_coauthor(self):
        """Добавление соавтора"""
        if not self.current_record_id:
            return

        name = self.coauthor_entry.get().strip()
        if not name:
            messagebox.showerror("Ошибка", "Введите имя соавтора")
            self.coauthor_entry.focus()
            return

        # Проверяем, нет ли уже такого соавтора
        coauthors = self.db.get_coauthors(self.current_record_id)
        if name in coauthors:
            messagebox.showwarning("Предупреждение", f"Соавтор '{name}' уже добавлен")
            self.coauthor_entry.delete(0, tk.END)
            return

        # Добавляем соавтора
        success = self.db.add_coauthor(self.current_record_id, name)

        if success:
            messagebox.showinfo("Успех", f"Соавтор {name} успешно добавлен")
            self.coauthor_entry.delete(0, tk.END)

            # Обновляем список соавторов
            self.load_coauthors(self.current_record_id)

            self.update_status(f"Добавлен соавтор: {name}")
        else:
            messagebox.showerror("Ошибка", "Не удалось добавить соавтора")

    def remove_selected_coauthor(self):
        """Удаление выбранного соавтора"""
        if not self.current_record_id:
            return

        selection = self.coauthors_listbox.curselection()
        if not selection:
            return

        # Получаем имя соавтора
        name = self.coauthors_listbox.get(selection[0])

        # Подтверждение
        if not messagebox.askyesno("Подтверждение", f"Удалить соавтора: {name}?"):
            return

        # Удаляем соавтора
        success = self.db.delete_coauthor(self.current_record_id, name)

        if success:
            messagebox.showinfo("Успех", f"Соавтор {name} удален")

            # Обновляем список соавторов
            self.load_coauthors(self.current_record_id)

            self.update_status(f"Удален соавтор: {name}")
        else:
            messagebox.showerror("Ошибка", "Не удалось удалить соавтора")

    def open_description(self):
        """Открытие описания в браузере (HTML preview)"""
        if not self.current_record_id:
            messagebox.showwarning("Предупреждение", "Сначала выберите запись")
            return

        # Получаем запись
        record = self.db.get_record_by_id(self.current_record_id)
        if not record:
            return

        # Конвертируем Markdown в HTML
        description = record.get('description', '')
        if not description:
            messagebox.showinfo("Информация", "У записи нет описания")
            return

        html_content = markdown.markdown(description, extensions=['fenced_code', 'tables'])

        # Создаем HTML файл
        os.makedirs('temp', exist_ok=True)
        html_path = os.path.join('temp', 'preview.html')

        with open(html_path, 'w', encoding='utf-8') as f:
            f.write(f'''
            <!DOCTYPE html>
            <html>
            <head>
                <meta charset="utf-8">
                <title>{record['title']}</title>
                <style>
                    body {{ 
                        font-family: Arial, sans-serif; 
                        margin: 40px; 
                        line-height: 1.6;
                        max-width: 800px;
                        margin: 40px auto;
                        padding: 20px;
                    }}
                    h1 {{ color: #333; border-bottom: 2px solid #eee; }}
                    pre {{ 
                        background-color: #f5f5f5; 
                        padding: 15px; 
                        border-radius: 5px;
                        overflow-x: auto;
                    }}
                    code {{ background-color: #f5f5f5; padding: 2px 4px; border-radius: 3px; }}
                    blockquote {{ 
                        border-left: 4px solid #ccc; 
                        padding-left: 15px; 
                        margin-left: 0;
                        color: #666;
                    }}
                    table {{ border-collapse: collapse; width: 100%; margin: 20px 0; }}
                    th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                    th {{ background-color: #f2f2f2; }}
                </style>
            </head>
            <body>
                <h1>{record['title']}</h1>
                <div style="color: #666; margin-bottom: 30px;">
                    Тип: {record['type']} | Год: {record['year']}
                </div>
                <hr>
                {html_content}
            </body>
            </html>
            ''')

        # Открываем в браузере
        try:
            webbrowser.open(f'file://{os.path.abspath(html_path)}')
            self.update_status("Описание открыто в браузере")
        except:
            messagebox.showerror("Ошибка", "Не удалось открыть браузер")

    def export_to_excel(self):
        """Экспорт всех записей в Excel"""
        records = self.db.get_all_records()
        if not records:
            messagebox.showinfo("Информация", "Нет записей для экспорта")
            return

        # Выбираем путь для сохранения
        filepath = filedialog.asksaveasfilename(
            defaultextension=".xlsx",
            filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")],
            initialfile="portfolio_export.xlsx"
        )

        if not filepath:
            return

        try:
            # Создаем DataFrame
            data = []
            for record in records:
                created_at = record['created_at']
                if hasattr(created_at, 'strftime'):
                    created_str = created_at.strftime('%Y-%m-%d %H:%M')
                else:
                    created_str = str(created_at)[:16]

                data.append({
                    'ID': record['id'],
                    'Название': record['title'],
                    'Тип': record['type'],
                    'Год': record['year'],
                    'Дата создания': created_str,
                    'Описание': record['description'][:100] + '...' if record['description'] and len(
                        record['description']) > 100 else record['description'] or ''
                })

            df = pd.DataFrame(data)

            # Сохраняем в Excel
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                df.to_excel(writer, sheet_name='Записи', index=False)

                # Добавляем лист с соавторами
                all_coauthors = []
                for record in records:
                    coauthors = self.db.get_coauthors(record['id'])
                    for coauthor in coauthors:
                        all_coauthors.append({
                            'ID записи': record['id'],
                            'Название': record['title'],
                            'Соавтор': coauthor
                        })

                if all_coauthors:
                    coauthors_df = pd.DataFrame(all_coauthors)
                    coauthors_df.to_excel(writer, sheet_name='Соавторы', index=False)

            messagebox.showinfo("Успех", f"Данные успешно экспортированы в:\n{filepath}")
            self.update_status(f"Экспорт в Excel выполнен: {filepath}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при экспорте: {str(e)}")
            print(traceback.format_exc())

    def export_to_word(self):
        """Экспорт выбранной записи в Word"""
        if not self.current_record_id:
            messagebox.showwarning("Предупреждение", "Сначала выберите запись")
            return

        record = self.db.get_record_by_id(self.current_record_id)
        if not record:
            return

        # Выбираем путь для сохранения
        filepath = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word files", "*.docx"), ("All files", "*.*")],
            initialfile=f"{record['title'].replace(' ', '_')}.docx"
        )

        if not filepath:
            return

        try:
            doc = Document()

            # Настраиваем стили
            self.setup_word_styles(doc)

            # Заголовок
            title = doc.add_heading(record['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Метаданные
            meta_para = doc.add_paragraph()
            meta_para.add_run("Тип: ").bold = True
            meta_para.add_run(record['type'])

            meta_para = doc.add_paragraph()
            meta_para.add_run("Год: ").bold = True
            meta_para.add_run(str(record['year']))

            created_at = record['created_at']
            if hasattr(created_at, 'strftime'):
                created_str = created_at.strftime('%d.%m.%Y %H:%M')
            else:
                created_str = str(created_at)[:16]

            meta_para = doc.add_paragraph()
            meta_para.add_run("Дата создания: ").bold = True
            meta_para.add_run(created_str)

            # Соавторы
            coauthors = self.db.get_coauthors(self.current_record_id)
            if coauthors:
                meta_para = doc.add_paragraph()
                meta_para.add_run("Соавторы: ").bold = True
                meta_para.add_run(", ".join(coauthors))

            doc.add_paragraph()

            # Описание
            if record.get('description'):
                doc.add_heading('Описание', level=1)

                # Простая обработка Markdown
                lines = record['description'].split('\n')
                for line in lines:
                    line = line.rstrip()
                    if not line:
                        doc.add_paragraph()
                    elif line.startswith('# '):
                        doc.add_heading(line[2:], level=1)
                    elif line.startswith('## '):
                        doc.add_heading(line[3:], level=2)
                    elif line.startswith('### '):
                        doc.add_heading(line[4:], level=3)
                    elif line.startswith('> '):
                        p = doc.add_paragraph(style='Intense Quote')
                        p.add_run(line[2:])
                    elif line.strip().startswith('- ') or line.strip().startswith('* '):
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(line.strip()[2:])
                    elif line.strip().startswith('1. '):
                        p = doc.add_paragraph(style='List Number')
                        p.add_run(line.strip()[3:])
                    elif '```' in line:
                        continue  # Пропускаем строки с кодом для простоты
                    else:
                        doc.add_paragraph(line)

            # Сохраняем документ
            doc.save(filepath)

            messagebox.showinfo("Успех", f"Запись успешно экспортирована в:\n{filepath}")
            self.update_status(f"Экспорт в Word выполнен: {filepath}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при экспорте: {str(e)}")
            print(traceback.format_exc())

    def setup_word_styles(self, doc):
        """Настройка стилей для Word документа"""
        # Основной стиль
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        # Настройка межстрочного интервала
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # Стиль для заголовков
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.font.name = 'Times New Roman'
            heading_style.font.bold = True

    def open_analytics(self):
        """Открытие окна аналитики"""
        # Создаем новое окно
        analytics_window = tk.Toplevel(self.root)
        analytics_window.title("Аналитика и отчётность")
        analytics_window.geometry("900x700")
        analytics_window.transient(self.root)
        analytics_window.grab_set()

        # Основной фрейм
        main_frame = ttk.Frame(analytics_window, padding="20")
        main_frame.pack(fill=tk.BOTH, expand=True)

        # Заголовок
        ttk.Label(main_frame, text="Аналитика портфолио",
                  font=("Arial", 16, "bold")).pack(pady=(0, 20))

        # Кнопка формирования отчёта
        ttk.Button(main_frame, text="Сформировать отчёт",
                   command=lambda: self.generate_report(analytics_window),
                   style="Accent.TButton").pack(pady=(0, 20))

        # Область для статистики
        stats_frame = ttk.LabelFrame(main_frame, text="Статистика", padding="10")
        stats_frame.pack(fill=tk.BOTH, expand=True, pady=(0, 20))

        # Получаем статистику
        stats = self.db.get_statistics()

        # Отображаем статистику
        stats_text = self.format_statistics(stats)
        stats_label = ttk.Label(stats_frame, text=stats_text, justify=tk.LEFT)
        stats_label.pack(anchor=tk.W, fill=tk.BOTH, expand=True)

        # Область для графиков
        self.create_charts(main_frame, stats)

    def format_statistics(self, stats):
        """Форматирование статистики для отображения"""
        text = ""

        text += f"Всего записей: {stats.get('total_records', 0)}\n"
        text += f"Уникальных соавторов: {stats.get('unique_coauthors', 0)}\n\n"

        text += "Распределение по типам:\n"
        for type_name, count in stats.get('type_distribution', {}).items():
            text += f"  • {type_name}: {count}\n"

        text += "\nРаспределение по годам:\n"
        for year, count in sorted(stats.get('year_distribution', {}).items()):
            text += f"  • {year}: {count}\n"

        text += "\nАктивность за 12 месяцев:\n"
        for month, count in sorted(stats.get('monthly_activity', {}).items()):
            text += f"  • {month}: {count}\n"

        return text

    def create_charts(self, parent, stats):
        """Создание графиков"""
        if not stats.get('total_records', 0):
            ttk.Label(parent, text="Нет данных для построения графиков").pack()
            return

        # Создаем фрейм для графиков
        charts_frame = ttk.Frame(parent)
        charts_frame.pack(fill=tk.BOTH, expand=True)

        try:
            # Создаем фигуру с графиками
            fig = Figure(figsize=(10, 8), dpi=100)

            # График 1: Распределение по типам
            if stats.get('type_distribution'):
                ax1 = fig.add_subplot(221)
                types = list(stats['type_distribution'].keys())
                counts = list(stats['type_distribution'].values())
                bars = ax1.bar(types, counts)
                ax1.set_title('Распределение по типам')
                ax1.set_xlabel('Тип записи')
                ax1.set_ylabel('Количество')
                ax1.tick_params(axis='x', rotation=45)

                # Добавляем значения на столбцы
                for bar, count in zip(bars, counts):
                    height = bar.get_height()
                    ax1.text(bar.get_x() + bar.get_width() / 2., height + 0.1,
                             f'{count}', ha='center', va='bottom', fontsize=9)

            # График 2: Распределение по годам
            if stats.get('year_distribution'):
                ax2 = fig.add_subplot(222)
                years = [str(year) for year in sorted(stats['year_distribution'].keys())]
                counts = [stats['year_distribution'][int(year)] for year in years]
                bars = ax2.bar(years, counts)
                ax2.set_title('Распределение по годам')
                ax2.set_xlabel('Год')
                ax2.set_ylabel('Количество')
                ax2.tick_params(axis='x', rotation=45)

            # График 3: Активность по месяцам
            if stats.get('monthly_activity'):
                ax3 = fig.add_subplot(223)
                months = list(stats['monthly_activity'].keys())
                counts = list(stats['monthly_activity'].values())
                ax3.plot(months, counts, marker='o', linestyle='-')
                ax3.set_title('Активность за 12 месяцев')
                ax3.set_xlabel('Месяц')
                ax3.set_ylabel('Количество записей')
                ax3.tick_params(axis='x', rotation=45)
                ax3.grid(True, alpha=0.3)

            # График 4: Круговая диаграмма типов
            if stats.get('type_distribution') and len(stats['type_distribution']) > 1:
                ax4 = fig.add_subplot(224)
                types = list(stats['type_distribution'].keys())
                counts = list(stats['type_distribution'].values())
                ax4.pie(counts, labels=types, autopct='%1.1f%%', startangle=90)
                ax4.set_title('Доля типов записей')
                ax4.axis('equal')

            fig.tight_layout()

            # Встраиваем график в Tkinter
            canvas = FigureCanvasTkAgg(fig, charts_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

        except Exception as e:
            print(f"Ошибка создания графиков: {e}")
            ttk.Label(charts_frame, text=f"Ошибка создания графиков: {e}").pack()

    def generate_report(self, analytics_window=None):
        """Формирование комплексного отчёта"""
        try:
            # Создаем директории
            os.makedirs('reports', exist_ok=True)
            os.makedirs('temp', exist_ok=True)

            # Получаем статистику
            stats = self.db.get_statistics()

            # Получаем последние 5 записей
            records = self.db.get_all_records()[:5]

            # 1. Экспорт в Excel
            excel_path = os.path.join('reports', 'portfolio_report.xlsx')
            success_excel = self.generate_excel_report(stats, records, excel_path)

            # 2. Генерация Word документа
            word_path = os.path.join('reports', 'portfolio_report.docx')
            success_word = self.generate_word_report(stats, records, word_path)

            if success_excel and success_word:
                messagebox.showinfo("Успех",
                                    f"Отчёт успешно сформирован!\n\n"
                                    f"Excel: {excel_path}\n"
                                    f"Word: {word_path}")

                if analytics_window:
                    analytics_window.destroy()

                self.update_status("Отчёт сформирован")
            else:
                messagebox.showwarning("Предупреждение",
                                       "Отчёт сформирован частично. Проверьте наличие файлов.")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при формировании отчёта:\n{str(e)}")
            print(traceback.format_exc())

    def generate_excel_report(self, stats, records, filepath):
        """Генерация Excel отчёта"""
        try:
            wb = Workbook()

            # Лист "Статистика"
            ws_stats = wb.active
            ws_stats.title = "Статистика"

            # Заголовок
            ws_stats.append(["Отчёт по портфолио"])
            ws_stats.append([f"Дата формирования: {datetime.datetime.now().strftime('%d.%m.%Y %H:%M:%S')}"])
            ws_stats.append([])

            # Ключевые показатели
            ws_stats.append(["Ключевые показатели"])
            ws_stats.append(["Всего записей:", stats.get('total_records', 0)])
            ws_stats.append(["Уникальных соавторов:", stats.get('unique_coauthors', 0)])
            ws_stats.append([])

            # Распределение по типам
            ws_stats.append(["Распределение по типам"])
            ws_stats.append(["Тип", "Количество"])
            for type_name, count in stats.get('type_distribution', {}).items():
                ws_stats.append([type_name, count])
            ws_stats.append([])

            # Распределение по годам
            ws_stats.append(["Распределение по годам"])
            ws_stats.append(["Год", "Количество"])
            for year, count in sorted(stats.get('year_distribution', {}).items()):
                ws_stats.append([year, count])
            ws_stats.append([])

            # Активность по месяцам
            ws_stats.append(["Активность за 12 месяцев"])
            ws_stats.append(["Месяц", "Количество"])
            for month, count in sorted(stats.get('monthly_activity', {}).items()):
                ws_stats.append([month, count])

            # Лист "Графики"
            ws_charts = wb.create_sheet("Графики")

            # Создаем и сохраняем графики
            self.create_excel_charts(wb, ws_charts, stats)

            # Лист "Последние записи"
            ws_records = wb.create_sheet("Последние записи")
            ws_records.append(["Последние 5 записей"])
            ws_records.append([])
            ws_records.append(["ID", "Название", "Тип", "Год", "Дата создания"])

            for record in records:
                created_at = record['created_at']
                if hasattr(created_at, 'strftime'):
                    created_str = created_at.strftime('%Y-%m-%d %H:%M')
                else:
                    created_str = str(created_at)[:16]

                ws_records.append([
                    record['id'],
                    record['title'],
                    record['type'],
                    record['year'],
                    created_str
                ])

            # Сохраняем файл
            wb.save(filepath)
            return True

        except Exception as e:
            print(f"Ошибка генерации Excel отчёта: {e}")
            traceback.print_exc()
            return False

    def create_excel_charts(self, wb, worksheet, stats):
        """Создание графиков для Excel"""
        try:
            # Создаем лист для данных графиков
            ws_data = wb.create_sheet("Данные_графиков")

            # Данные для графиков распределения по типам
            ws_data.append(["Тип", "Количество"])
            for type_name, count in stats.get('type_distribution', {}).items():
                ws_data.append([type_name, count])

            ws_data.append([])
            ws_data.append(["Год", "Количество"])
            for year, count in sorted(stats.get('year_distribution', {}).items()):
                ws_data.append([year, count])

            # Создаем столбчатую диаграмму для типов
            if stats.get('type_distribution'):
                chart1 = BarChart()
                chart1.type = "col"
                chart1.style = 10
                chart1.title = "Распределение записей по типам"
                chart1.y_axis.title = 'Количество'
                chart1.x_axis.title = 'Тип'

                data = Reference(ws_data, min_col=2, min_row=1, max_row=len(stats['type_distribution']) + 1)
                cats = Reference(ws_data, min_col=1, min_row=2, max_row=len(stats['type_distribution']) + 1)
                chart1.add_data(data, titles_from_data=True)
                chart1.set_categories(cats)
                chart1.shape = 4
                worksheet.add_chart(chart1, "A1")

            # Создаем столбчатую диаграмму для годов
            if stats.get('year_distribution'):
                chart2 = BarChart()
                chart2.type = "col"
                chart2.style = 10
                chart2.title = "Распределение записей по годам"
                chart2.y_axis.title = 'Количество'
                chart2.x_axis.title = 'Год'

                start_row = len(stats['type_distribution']) + 4
                end_row = start_row + len(stats['year_distribution'])
                data = Reference(ws_data, min_col=2, min_row=start_row, max_row=end_row)
                cats = Reference(ws_data, min_col=1, min_row=start_row + 1, max_row=end_row)
                chart2.add_data(data, titles_from_data=True)
                chart2.set_categories(cats)
                chart2.shape = 4
                worksheet.add_chart(chart2, "A20")

            # Скрываем лист с данными
            ws_data.sheet_state = 'hidden'

        except Exception as e:
            print(f"Ошибка создания графиков Excel: {e}")

    def generate_word_report(self, stats, records, filepath):
        """Генерация Word отчёта"""
        try:
            doc = Document()

            # Настройка стилей
            self.setup_word_report_styles(doc)

            # Титульный лист
            self.add_title_page(doc)

            # Содержание
            doc.add_page_break()
            self.add_table_of_contents(doc)

            # Ключевые показатели
            doc.add_page_break()
            self.add_key_metrics(doc, stats)

            # Статистика
            self.add_statistics(doc, stats)

            # Графики
            self.add_charts_to_word(doc, stats)

            # Последние записи
            self.add_recent_records(doc, records)

            # Сохраняем документ
            doc.save(filepath)
            return True

        except Exception as e:
            print(f"Ошибка генерации Word отчёта: {e}")
            traceback.print_exc()
            return False

    def setup_word_report_styles(self, doc):
        """Настройка стилей для отчёта Word"""
        # Основной стиль
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)

        # Межстрочный интервал
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        # Стили заголовков
        for i in range(1, 4):
            heading = doc.styles[f'Heading {i}']
            heading.font.name = 'Times New Roman'
            heading.font.bold = True

        # Стиль для таблиц
        try:
            table_style = doc.styles.add_style('ReportTable', WD_STYLE_TYPE.PARAGRAPH)
            table_style.font.name = 'Times New Roman'
            table_style.font.size = Pt(11)
        except:
            pass

    def add_title_page(self, doc):
        """Добавление титульного листа"""
        # Пустые строки для центрирования
        for _ in range(10):
            doc.add_paragraph()

        # Название отчёта
        title = doc.add_heading('Отчёт по портфолио', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Дата формирования
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.add_run(f'Дата формирования: {datetime.datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')

        # Пустые строки в конце
        for _ in range(10):
            doc.add_paragraph()

    def add_table_of_contents(self, doc):
        """Добавление содержания"""
        heading = doc.add_heading('Содержание', 1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("1. Ключевые показатели")
        doc.add_paragraph("2. Статистика")
        doc.add_paragraph("3. Графики и диаграммы")
        doc.add_paragraph("4. Последние записи")

    def add_key_metrics(self, doc, stats):
        """Добавление ключевых показателей"""
        heading = doc.add_heading('1. Ключевые показатели', 1)

        table = doc.add_table(rows=3, cols=2)
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Заполняем таблицу
        table.cell(0, 0).text = 'Всего записей'
        table.cell(0, 1).text = str(stats.get('total_records', 0))

        table.cell(1, 0).text = 'Уникальных соавторов'
        table.cell(1, 1).text = str(stats.get('unique_coauthors', 0))

        table.cell(2, 0).text = 'Дата формирования отчёта'
        table.cell(2, 1).text = datetime.datetime.now().strftime('%d.%m.%Y %H:%M:%S')

    def add_statistics(self, doc, stats):
        """Добавление статистики"""
        doc.add_heading('2. Статистика', 1)

        # Распределение по типам
        doc.add_heading('Распределение по типам', 2)
        if stats.get('type_distribution'):
            table = doc.add_table(rows=len(stats['type_distribution']) + 1, cols=2)
            table.style = 'Light Grid Accent 1'

            table.cell(0, 0).text = 'Тип записи'
            table.cell(0, 1).text = 'Количество'

            for i, (type_name, count) in enumerate(stats['type_distribution'].items(), 1):
                table.cell(i, 0).text = type_name
                table.cell(i, 1).text = str(count)

        # Распределение по годам
        doc.add_heading('Распределение по годам', 2)
        if stats.get('year_distribution'):
            table = doc.add_table(rows=len(stats['year_distribution']) + 1, cols=2)
            table.style = 'Light Grid Accent 1'

            table.cell(0, 0).text = 'Год'
            table.cell(0, 1).text = 'Количество'

            for i, (year, count) in enumerate(sorted(stats['year_distribution'].items()), 1):
                table.cell(i, 0).text = str(year)
                table.cell(i, 1).text = str(count)

    def add_charts_to_word(self, doc, stats):
        """Добавление графиков в Word"""
        doc.add_heading('3. Графики и диаграммы', 1)

        # Создаем и сохраняем графики
        os.makedirs('temp', exist_ok=True)

        # График распределения по типам
        if stats.get('type_distribution'):
            plt.figure(figsize=(10, 6))
            types = list(stats['type_distribution'].keys())
            counts = list(stats['type_distribution'].values())

            bars = plt.bar(types, counts)
            plt.title('Распределение записей по типам')
            plt.xlabel('Тип записи')
            plt.ylabel('Количество')
            plt.xticks(rotation=45)
            plt.tight_layout()

            # Добавляем значения на столбцы
            for bar, count in zip(bars, counts):
                height = bar.get_height()
                plt.text(bar.get_x() + bar.get_width() / 2., height + 0.1,
                         f'{count}', ha='center', va='bottom')

            chart_path = os.path.join('temp', 'type_chart.png')
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()

            # Вставляем график в документ
            doc.add_paragraph("Распределение записей по типам:")
            doc.add_picture(chart_path, width=Inches(6))

        # График распределения по годам
        if stats.get('year_distribution'):
            plt.figure(figsize=(10, 6))
            years = [str(year) for year in sorted(stats['year_distribution'].keys())]
            counts = [stats['year_distribution'][int(year)] for year in years]

            bars = plt.bar(years, counts)
            plt.title('Распределение записей по годам')
            plt.xlabel('Год')
            plt.ylabel('Количество')
            plt.xticks(rotation=45)
            plt.tight_layout()

            chart_path = os.path.join('temp', 'year_chart.png')
            plt.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close()

            # Вставляем график в документ
            doc.add_paragraph("Распределение записей по годам:")
            doc.add_picture(chart_path, width=Inches(6))

    def add_recent_records(self, doc, records):
        """Добавление последних записей"""
        doc.add_heading('4. Последние записи', 1)

        if records:
            table = doc.add_table(rows=len(records) + 1, cols=4)
            table.style = 'Light Grid Accent 1'

            # Заголовки
            headers = ['Название', 'Тип', 'Год', 'Дата создания']
            for i, header in enumerate(headers):
                table.cell(0, i).text = header
                table.cell(0, i).paragraphs[0].runs[0].bold = True

            # Данные
            for i, record in enumerate(records, 1):
                created_at = record['created_at']
                if hasattr(created_at, 'strftime'):
                    created_str = created_at.strftime('%d.%m.%Y %H:%M')
                else:
                    created_str = str(created_at)[:16]

                table.cell(i, 0).text = record['title']
                table.cell(i, 1).text = record['type']
                table.cell(i, 2).text = str(record['year'])
                table.cell(i, 3).text = created_str
        else:
            doc.add_paragraph("Нет записей для отображения")

    def clear_form(self):
        """Очистка формы"""
        self.current_record_id = None

        self.title_entry.delete(0, tk.END)
        self.type_combobox.set("Статья")
        self.year_spinbox.set(datetime.datetime.now().year)
        self.text_editor.delete(1.0, tk.END)
        self.coauthor_entry.delete(0, tk.END)

        # Очищаем список соавторов
        self.coauthors_listbox.delete(0, tk.END)

        # Отключаем кнопки
        self.save_btn.config(state=tk.DISABLED)
        self.delete_btn.config(state=tk.DISABLED)
        self.add_coauthor_btn.config(state=tk.DISABLED)
        self.remove_coauthor_btn.config(state=tk.DISABLED)

    def update_status(self, message):
        """Обновление статус бара"""
        timestamp = datetime.datetime.now().strftime("%H:%M:%S")
        self.status_bar.config(text=f"[{timestamp}] {message}")


def main():
    """Основная функция приложения"""
    try:
        root = tk.Tk()

        # Настройка стиля
        style = ttk.Style()
        style.theme_use('clam')

        # Создание приложения
        app = PortfolioApp(root)

        # Запуск главного цикла
        root.mainloop()

    except Exception as e:
        print(f"Критическая ошибка: {e}")
        traceback.print_exc()
        messagebox.showerror("Ошибка", f"Не удалось запустить приложение:\n{str(e)}")


if __name__ == "__main__":
    # Устанавливаем кодировку
    if sys.platform == 'win32':
        import locale

        locale.setlocale(locale.LC_ALL, '')

    main()
