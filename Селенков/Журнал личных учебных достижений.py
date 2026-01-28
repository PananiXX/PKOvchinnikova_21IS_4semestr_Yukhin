# -*- coding: utf-8 -*-
"""
Журнал личных учебных достижений
Используется Python 3.8+, tkinter, SQLite, python-docx
"""

import tkinter as tk
from tkinter import ttk, messagebox
import sqlite3
import json
import os
from datetime import datetime
from docx import Document


class AchievementJournal:
    def __init__(self, root):
        self.root = root
        self.root.title("Журнал достижений")
        self.root.geometry("600x450")  # Немного увеличено для лучшего отображения

        # Загрузка типов из JSON
        self.types_list = self.load_types()

        # Создание вкладок
        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill="both", expand=True)

        self.tab_add = ttk.Frame(self.notebook)
        self.tab_list = ttk.Frame(self.notebook)

        self.notebook.add(self.tab_add, text="Добавить")
        self.notebook.add(self.tab_list, text="Мои достижения")

        # Инициализация БД (SQLite как требует ТЗ)
        self.init_db()

        # Создание форм
        self.create_add_form()
        self.create_list_form()

    def load_types(self):
        """Загрузка типов достижений из JSON-файла"""
        try:
            with open("types.json", "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, list):
                    return data
                else:
                    print("Ошибка: types.json должен содержать список")
                    return ["Олимпиада", "Сертификат", "Проект", "Экзамен", "Конференция"]
        except FileNotFoundError:
            print("Файл types.json не найден, используются типы по умолчанию")
            return ["Олимпиада", "Сертификат", "Проект", "Экзамен", "Конференция"]
        except json.JSONDecodeError:
            print("Ошибка чтения types.json, используются типы по умолчанию")
            return ["Олимпиада", "Сертификат", "Проект", "Экзамен", "Конференция"]

    def init_db(self):
        """Инициализация базы данных SQLite (как требует ТЗ)"""
        try:
            conn = sqlite3.connect("достижения.db")
            cursor = conn.cursor()

            # Создание таблицы, если её нет
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS достижения (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    название TEXT NOT NULL,
                    дата TEXT NOT NULL,
                    тип TEXT NOT NULL,
                    уровень TEXT NOT NULL,
                    описание TEXT
                )
            """)

            conn.commit()
            conn.close()
            print("База данных SQLite инициализирована успешно")
        except Exception as e:
            print(f"Ошибка инициализации БД: {e}")

    def create_add_form(self):
        """Создание формы для добавления достижений"""
        panel = self.tab_add

        # Название
        tk.Label(panel, text="Название:").pack(anchor="w", padx=10, pady=5)
        self.name_entry = tk.Entry(panel, width=60)
        self.name_entry.pack(padx=10, pady=5)

        # Дата
        tk.Label(panel, text="Дата (ГГГГ-ММ-ДД):").pack(anchor="w", padx=10, pady=5)
        self.date_entry = tk.Entry(panel, width=60)
        self.date_entry.pack(padx=10, pady=5)

        # Тип
        tk.Label(panel, text="Тип:").pack(anchor="w", padx=10, pady=5)
        self.type_combobox = ttk.Combobox(panel, values=self.types_list, state="readonly", width=57)
        self.type_combobox.pack(padx=10, pady=5)
        if self.types_list:
            self.type_combobox.current(0)

        # Уровень
        tk.Label(panel, text="Уровень:").pack(anchor="w", padx=10, pady=5)
        self.level_combobox = ttk.Combobox(
            panel,
            values=["локальный", "региональный", "национальный", "международный"],
            state="readonly",
            width=57
        )
        self.level_combobox.pack(padx=10, pady=5)
        self.level_combobox.current(0)

        # Описание
        tk.Label(panel, text="Описание:").pack(anchor="w", padx=10, pady=5)
        self.desc_text = tk.Text(panel, height=6, width=60)
        self.desc_text.pack(padx=10, pady=5)

        # Кнопка сохранения
        self.save_btn = tk.Button(panel, text="Сохранить", command=self.on_save, bg="lightblue")
        self.save_btn.pack(padx=10, pady=15)

    def validate_date(self, date_text):
        """Проверка корректности формата даты"""
        try:
            datetime.strptime(date_text, "%Y-%m-%d")
            return True
        except ValueError:
            return False

    def on_save(self):
        """Обработчик сохранения достижения"""
        name = self.name_entry.get().strip()
        date = self.date_entry.get().strip()
        typ = self.type_combobox.get()
        level = self.level_combobox.get()
        desc = self.desc_text.get("1.0", "end-1c").strip()

        # Валидация
        if not name:
            messagebox.showerror("Ошибка", "Введите название достижения")
            return

        if not date:
            messagebox.showerror("Ошибка", "Введите дату")
            return

        if not self.validate_date(date):
            messagebox.showerror("Ошибка", "Неверный формат даты. Используйте ГГГГ-ММ-ДД")
            return

        # Сохранение в базу данных
        if self.save_to_db(name, date, typ, level, desc):
            messagebox.showinfo("Успех", f"Достижение '{name}' сохранено!")
            self.clear_form()
            self.refresh_list()  # Обновляем список сразу после сохранения

    def save_to_db(self, name, date, typ, level, desc):
        """Сохранение данных в SQLite"""
        try:
            conn = sqlite3.connect("достижения.db")
            cursor = conn.cursor()
            cursor.execute(
                "INSERT INTO достижения (название, дата, тип, уровень, описание) VALUES (?, ?, ?, ?, ?)",
                (name, date, typ, level, desc)
            )
            conn.commit()
            conn.close()
            print(f"Сохранено в БД: {name}")
            return True
        except Exception as e:
            print(f"Ошибка сохранения в БД: {e}")
            messagebox.showerror("Ошибка", f"Ошибка сохранения: {e}")
            return False

    def clear_form(self):
        """Очистка формы после сохранения"""
        self.name_entry.delete(0, tk.END)
        self.date_entry.delete(0, tk.END)
        self.desc_text.delete("1.0", tk.END)
        if self.types_list:
            self.type_combobox.current(0)
        self.level_combobox.current(0)

    def create_list_form(self):
        """Создание формы для отображения списка достижений"""
        # Список достижений с полосой прокрутки
        frame = tk.Frame(self.tab_list)
        frame.pack(fill="both", expand=True, padx=10, pady=10)

        # Полоса прокрутки
        scrollbar = tk.Scrollbar(frame)
        scrollbar.pack(side="right", fill="y")

        # Список
        self.listbox = tk.Listbox(frame, width=80, height=15, yscrollcommand=scrollbar.set)
        self.listbox.pack(side="left", fill="both", expand=True)

        scrollbar.config(command=self.listbox.yview)

        # Фрейм для кнопок
        button_frame = tk.Frame(self.tab_list)
        button_frame.pack(pady=10)

        # Кнопка обновления
        refresh_btn = tk.Button(button_frame, text="Обновить список", command=self.refresh_list, bg="lightgreen")
        refresh_btn.pack(side="left", padx=5)

        # Кнопка экспорта
        export_btn = tk.Button(button_frame, text="Экспорт в Word", command=self.export_to_word, bg="lightyellow")
        export_btn.pack(side="left", padx=5)

        # Кнопка удаления
        delete_btn = tk.Button(button_frame, text="Удалить выбранное", command=self.delete_selected, bg="lightcoral")
        delete_btn.pack(side="left", padx=5)

        # Первоначальное обновление списка
        self.refresh_list()

    def load_records(self):
        """Загрузка записей из SQLite"""
        try:
            conn = sqlite3.connect("достижения.db")
            cursor = conn.cursor()
            cursor.execute("SELECT дата, название, тип, уровень FROM достижения ORDER BY дата DESC")
            rows = cursor.fetchall()
            conn.close()
            return rows
        except Exception as e:
            print(f"Ошибка загрузки данных: {e}")
            return []

    def load_records_with_desc(self):
        """Загрузка записей с описанием из SQLite"""
        try:
            conn = sqlite3.connect("достижения.db")
            cursor = conn.cursor()
            cursor.execute("SELECT дата, название, тип, уровень, описание FROM достижения ORDER BY дата DESC")
            rows = cursor.fetchall()
            conn.close()
            return rows
        except Exception as e:
            print(f"Ошибка загрузки данных: {e}")
            return []

    def refresh_list(self):
        """Обновление списка достижений"""
        self.listbox.delete(0, tk.END)
        records = self.load_records()

        if not records:
            self.listbox.insert(tk.END, "Нет сохранённых достижений")
            return

        for date, name, typ, level in records:
            # Обрезаем длинные названия
            display_name = name[:50] + "..." if len(name) > 50 else name
            self.listbox.insert(tk.END, f"{date} - {display_name} ({typ}, {level})")

        # Сохраняем полные названия для доступа при удалении
        self.current_records = records

    def delete_selected(self):
        """Удаление выбранной записи"""
        selection = self.listbox.curselection()
        if not selection:
            messagebox.showwarning("Внимание", "Выберите запись для удаления")
            return

        index = selection[0]
        if self.listbox.get(index) == "Нет сохранённых достижений":
            return

        # Подтверждение удаления
        if not messagebox.askyesno("Подтверждение", "Удалить выбранное достижение?"):
            return

        # Получаем данные из текущих записей
        if hasattr(self, 'current_records') and self.current_records:
            try:
                date, name, _, _ = self.current_records[index]

                # Удаление из базы данных
                conn = sqlite3.connect("достижения.db")
                cursor = conn.cursor()
                cursor.execute("DELETE FROM достижения WHERE название = ? AND дата = ?", (name, date))
                conn.commit()
                conn.close()

                messagebox.showinfo("Успех", f"Достижение '{name[:30]}...' удалено")
                self.refresh_list()

            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка удаления: {e}")

    def export_to_word(self):
        """Экспорт достижений в документ Word"""
        records = self.load_records_with_desc()

        if not records:
            messagebox.showinfo("Информация", "Нет данных для экспорта")
            return

        try:
            doc = Document()
            doc.add_heading("Личные учебные достижения", 0)

            # Добавляем информацию о количестве
            doc.add_paragraph(f"Всего достижений: {len(records)}")
            doc.add_paragraph()

            for i, (date, name, typ, level, desc) in enumerate(records, 1):
                # Заголовок для каждого достижения
                doc.add_heading(f"{i}. {name}", level=2)

                # Информация о достижении
                info_para = doc.add_paragraph()
                info_para.add_run("Дата: ").bold = True
                info_para.add_run(f"{date}\n")

                info_para.add_run("Тип: ").bold = True
                info_para.add_run(f"{typ}\n")

                info_para.add_run("Уровень: ").bold = True
                info_para.add_run(f"{level}\n")

                # Описание, если есть
                if desc:
                    doc.add_heading("Описание:", level=3)
                    doc.add_paragraph(desc)

                doc.add_paragraph()  # Пустая строка между записями

            # Сохраняем документ
            filename = "достижения.docx"
            doc.save(filename)

            messagebox.showinfo("Успех", f"Отчёт сохранён в файл: {filename}")
            print(f"Отчёт экспортирован: {filename}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка экспорта в Word: {e}")
            print(f"Ошибка экспорта: {e}")


def main():
    """Основная функция запуска приложения"""
    root = tk.Tk()
    app = AchievementJournal(root)

    # Центрирование окна
    root.update_idletasks()
    width = root.winfo_width()
    height = root.winfo_height()
    x = (root.winfo_screenwidth() // 2) - (width // 2)
    y = (root.winfo_screenheight() // 2) - (height // 2)
    root.geometry(f'{width}x{height}+{x}+{y}')

    root.mainloop()


if __name__ == "__main__":

    main()
