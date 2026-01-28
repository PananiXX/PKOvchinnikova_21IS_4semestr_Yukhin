import tkinter as tk
from tkinter import ttk, messagebox
import psycopg2
from datetime import datetime
import json
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class PortfolioApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Портфолио исследователя")
        self.root.geometry("1000x750")

        self.entry_types = ["Проект", "Публикация", "Конференция", "Практика", "Грант"]
        self.current_specialty = None
        self.competencies = []
        self.used_keywords = []

        self.setup_database()
        self.load_specialties()
        self.load_used_keywords()

        self.notebook = ttk.Notebook(root)
        self.notebook.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.create_main_tab()
        self.create_research_map_tab()
        self.create_achievements_tab()
        self.create_competencies_tab()
        self.create_goals_tab()

        self.root.after(100, self.update_all_tabs)

    def update_all_tabs(self):
        self.update_research_map()
        self.update_achievements()
        self.update_competencies()
        self.update_goals()

    def setup_database(self):
        try:
            self.conn = psycopg2.connect(
                host="localhost",
                database="21ис2",
                user="postgres",
                password="1111",
                port="5432"
            )
            self.cursor = self.conn.cursor()
            self.create_tables()
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось подключиться к базе данных: {str(e)}")
            self.root.destroy()

    def create_tables(self):
        queries = [
            """
            CREATE TABLE IF NOT EXISTS entries (
                id SERIAL PRIMARY KEY,
                название TEXT NOT NULL,
                тип TEXT NOT NULL,
                дата DATE NOT NULL,
                описание TEXT,
                соавторы TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS keywords (
                id SERIAL PRIMARY KEY,
                keyword TEXT UNIQUE NOT NULL
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS entry_keywords (
                entry_id INTEGER REFERENCES entries(id) ON DELETE CASCADE,
                keyword_id INTEGER REFERENCES keywords(id) ON DELETE CASCADE,
                PRIMARY KEY (entry_id, keyword_id)
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS achievements (
                id SERIAL PRIMARY KEY,
                название TEXT UNIQUE NOT NULL,
                описание TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS user_achievements (
                user_id INTEGER DEFAULT 1,
                achievement_id INTEGER REFERENCES achievements(id) ON DELETE CASCADE,
                получено TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                PRIMARY KEY (user_id, achievement_id)
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS competencies (
                id SERIAL PRIMARY KEY,
                название TEXT NOT NULL,
                категория TEXT
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS entry_competencies (
                entry_id INTEGER REFERENCES entries(id) ON DELETE CASCADE,
                competency_id INTEGER REFERENCES competencies(id) ON DELETE CASCADE,
                уровень INTEGER CHECK (уровень >= 1 AND уровень <= 5),
                PRIMARY KEY (entry_id, competency_id)
            )
            """,
            """
            CREATE TABLE IF NOT EXISTS goals (
                id SERIAL PRIMARY KEY,
                описание TEXT NOT NULL,
                тип TEXT NOT NULL,
                цель_значение INTEGER,
                текущее_значение INTEGER DEFAULT 0,
                завершено BOOLEAN DEFAULT FALSE
            )
            """
        ]

        for query in queries:
            try:
                self.cursor.execute(query)
            except Exception as e:
                print(f"Ошибка при создании таблицы: {e}")

        self.cursor.execute("SELECT COUNT(*) FROM achievements")
        if self.cursor.fetchone()[0] == 0:
            achievements_data = [
                ("Первый шаг", "Создана первая запись"),
                ("Командный игрок", "Три и более записи с соавторами"),
                ("Разносторонний", "Записи минимум трёх разных типов"),
                ("Плодотворный год", "Три и более записи за один календарный год"),
                ("Словобог", "Суммарный объём описаний превысил 5000 символов")
            ]
            for name, desc in achievements_data:
                self.cursor.execute(
                    "INSERT INTO achievements (название, описание) VALUES (%s, %s) ON CONFLICT (название) DO NOTHING",
                    (name, desc))

        self.conn.commit()

    def load_used_keywords(self):
        try:
            self.cursor.execute("SELECT keyword FROM keywords ORDER BY keyword")
            self.used_keywords = [row[0] for row in self.cursor.fetchall()]
        except:
            self.used_keywords = []

    def load_specialties(self):
        try:
            with open("competencies.json", "r", encoding="utf-8") as f:
                self.specialties_data = json.load(f)
        except FileNotFoundError:
            self.specialties_data = {
                "Информационные системы": {
                    "competencies": [
                        {"название": "Программирование", "категория": "Технические"},
                        {"название": "Работа с БД", "категория": "Технические"},
                        {"название": "Проектирование систем", "категория": "Технические"},
                        {"название": "Презентация результатов", "категория": "Коммуникационные"},
                        {"название": "Научное письмо", "категория": "Коммуникационные"},
                        {"название": "Командная работа", "категория": "Социальные"},
                        {"название": "Управление проектами", "категория": "Организационные"}
                    ]
                }
            }
            with open("competencies.json", "w", encoding="utf-8") as f:
                json.dump(self.specialties_data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка загрузки компетенций: {str(e)}")
            self.specialties_data = {}

    def create_main_tab(self):
        main_tab = ttk.Frame(self.notebook)
        self.notebook.add(main_tab, text="Добавить запись")

        main_frame = ttk.Frame(main_tab)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        left_frame = ttk.LabelFrame(main_frame, text="Новая запись")
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5, pady=5)

        ttk.Label(left_frame, text="Название:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.title_entry = ttk.Entry(left_frame, width=40)
        self.title_entry.grid(row=0, column=1, padx=5, pady=5, sticky=tk.W)

        ttk.Label(left_frame, text="Тип:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        self.type_combo = ttk.Combobox(left_frame, values=self.entry_types, state="readonly", width=37)
        self.type_combo.grid(row=1, column=1, padx=5, pady=5, sticky=tk.W)

        ttk.Label(left_frame, text="Дата (ГГГГ-ММ-ДД):").grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        self.date_entry = ttk.Entry(left_frame, width=40)
        self.date_entry.grid(row=2, column=1, padx=5, pady=5, sticky=tk.W)
        self.date_entry.insert(0, datetime.now().strftime("%Y-%m-%d"))

        ttk.Label(left_frame, text="Описание:").grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
        self.desc_text = tk.Text(left_frame, width=40, height=5)
        self.desc_text.grid(row=3, column=1, padx=5, pady=5, sticky=tk.W)

        ttk.Label(left_frame, text="Соавторы (через запятую):").grid(row=4, column=0, sticky=tk.W, padx=5, pady=5)
        self.coauthors_entry = ttk.Entry(left_frame, width=40)
        self.coauthors_entry.grid(row=4, column=1, padx=5, pady=5, sticky=tk.W)

        ttk.Label(left_frame, text="Ключевые слова (до 5, через запятую):").grid(row=5, column=0, sticky=tk.W, padx=5,
                                                                                 pady=5)
        self.keywords_combo = ttk.Combobox(left_frame, width=38)
        self.keywords_combo.grid(row=5, column=1, padx=5, pady=5, sticky=tk.W)
        self.keywords_combo['values'] = self.used_keywords
        self.keywords_combo.bind('<KeyRelease>', self.update_keywords_suggestions)

        ttk.Label(left_frame, text="Компетенции (выберите 1-3):").grid(row=6, column=0, sticky=tk.W, padx=5, pady=5)
        self.competencies_frame = ttk.Frame(left_frame)
        self.competencies_frame.grid(row=6, column=1, padx=5, pady=5, sticky=tk.W)

        ttk.Label(left_frame, text="Уровень компетенций (1-5):").grid(row=7, column=0, sticky=tk.W, padx=5, pady=5)
        self.level_combo = ttk.Combobox(left_frame, values=[1, 2, 3, 4, 5], state="readonly", width=37)
        self.level_combo.grid(row=7, column=1, padx=5, pady=5, sticky=tk.W)
        self.level_combo.set(3)

        self.comp_vars = []
        self.create_competencies_widgets()

        ttk.Button(left_frame, text="Сохранить запись", command=self.save_entry).grid(row=8, column=0, columnspan=2,
                                                                                      pady=20)

        right_frame = ttk.LabelFrame(main_frame, text="Предыдущие ключевые слова")
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.keywords_listbox = tk.Listbox(right_frame)
        self.keywords_listbox.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        self.update_keywords_listbox()

        bottom_frame = ttk.Frame(main_tab)
        bottom_frame.pack(fill=tk.X, padx=10, pady=10)

        ttk.Button(bottom_frame, text="Экспорт в Word", command=self.export_to_word).pack(side=tk.LEFT, padx=5)
        ttk.Button(bottom_frame, text="Очистить форму", command=self.clear_form).pack(side=tk.LEFT, padx=5)

    def update_keywords_suggestions(self, event=None):
        current_text = self.keywords_combo.get()
        if current_text:
            suggestions = [kw for kw in self.used_keywords if current_text.lower() in kw.lower()]
            self.keywords_combo['values'] = suggestions
        else:
            self.keywords_combo['values'] = self.used_keywords

    def update_keywords_listbox(self):
        self.keywords_listbox.delete(0, tk.END)
        for keyword in self.used_keywords:
            self.keywords_listbox.insert(tk.END, keyword)

    def create_competencies_widgets(self):
        for widget in self.competencies_frame.winfo_children():
            widget.destroy()

        self.comp_vars = []
        self.comp_checkboxes = []

        if self.current_specialty and self.competencies:
            for i, comp in enumerate(self.competencies):
                var = tk.BooleanVar()
                cb = ttk.Checkbutton(self.competencies_frame, text=comp['название'], variable=var)
                cb.grid(row=i // 2, column=i % 2, sticky=tk.W, padx=5, pady=2)
                self.comp_vars.append((var, comp['id']))
                self.comp_checkboxes.append(cb)
        else:
            ttk.Label(self.competencies_frame, text="Выберите специальность во вкладке 'Компетенции'").pack()

    def save_entry(self):
        try:
            title = self.title_entry.get()
            entry_type = self.type_combo.get()
            date_str = self.date_entry.get()
            description = self.desc_text.get("1.0", tk.END).strip()
            coauthors = self.coauthors_entry.get()

            if not all([title, entry_type, date_str]):
                messagebox.showwarning("Предупреждение", "Заполните обязательные поля: Название, Тип, Дата")
                return

            try:
                datetime.strptime(date_str, "%Y-%m-%d")
            except ValueError:
                messagebox.showwarning("Предупреждение", "Некорректный формат даты. Используйте ГГГГ-ММ-ДД")
                return

            selected_comps = [(var[1], self.level_combo.get()) for var in self.comp_vars if var[0].get()]
            if len(selected_comps) == 0:
                messagebox.showwarning("Предупреждение", "Выберите хотя бы одну компетенцию")
                return
            elif len(selected_comps) > 3:
                messagebox.showwarning("Предупреждение", "Можно выбрать не более 3 компетенций")
                return

            keywords_input = self.keywords_combo.get()
            keywords = [k.strip() for k in keywords_input.split(",") if k.strip()]
            if len(keywords) > 5:
                keywords = keywords[:5]
                messagebox.showinfo("Информация", "Выбрано более 5 ключевых слов. Сохранены первые 5.")

            self.cursor.execute(
                "INSERT INTO entries (название, тип, дата, описание, соавторы) VALUES (%s, %s, %s, %s, %s) RETURNING id",
                (title, entry_type, date_str, description, coauthors)
            )
            entry_id = self.cursor.fetchone()[0]

            for keyword in keywords:
                self.cursor.execute(
                    "INSERT INTO keywords (keyword) VALUES (%s) ON CONFLICT (keyword) DO UPDATE SET keyword = EXCLUDED.keyword RETURNING id",
                    (keyword,))
                keyword_id = self.cursor.fetchone()[0]
                self.cursor.execute(
                    "INSERT INTO entry_keywords (entry_id, keyword_id) VALUES (%s, %s) ON CONFLICT DO NOTHING",
                    (entry_id, keyword_id))

            for comp_id, level in selected_comps:
                self.cursor.execute(
                    "INSERT INTO entry_competencies (entry_id, competency_id, уровень) VALUES (%s, %s, %s) ON CONFLICT DO NOTHING",
                    (entry_id, comp_id, level))

            self.conn.commit()

            self.load_used_keywords()
            self.update_keywords_listbox()
            self.keywords_combo['values'] = self.used_keywords

            self.check_achievements(entry_id)

            messagebox.showinfo("Успех", "Запись сохранена")

            self.update_all_tabs()

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при сохранении: {str(e)}")

    def clear_form(self):
        self.title_entry.delete(0, tk.END)
        self.type_combo.set('')
        self.date_entry.delete(0, tk.END)
        self.date_entry.insert(0, datetime.now().strftime("%Y-%m-%d"))
        self.desc_text.delete("1.0", tk.END)
        self.coauthors_entry.delete(0, tk.END)
        self.keywords_combo.set('')
        for var, _ in self.comp_vars:
            var.set(False)

    def check_achievements(self, entry_id):
        user_id = 1

        self.cursor.execute("SELECT COUNT(*) FROM entries")
        count = self.cursor.fetchone()[0]
        if count == 1:
            self.unlock_achievement("Первый шаг", user_id)

        self.cursor.execute(
            "SELECT COUNT(*) FROM entries WHERE соавторы IS NOT NULL AND соавторы != '' AND соавторы != ' '")
        coauthored = self.cursor.fetchone()[0]
        if coauthored >= 3:
            self.unlock_achievement("Командный игрок", user_id)

        self.cursor.execute("SELECT COUNT(DISTINCT тип) FROM entries")
        types_count = self.cursor.fetchone()[0]
        if types_count >= 3:
            self.unlock_achievement("Разносторонний", user_id)

        current_year = datetime.now().year
        self.cursor.execute("SELECT COUNT(*) FROM entries WHERE EXTRACT(YEAR FROM дата) = %s", (current_year,))
        year_count = self.cursor.fetchone()[0]
        if year_count >= 3:
            self.unlock_achievement("Плодотворный год", user_id)

        self.cursor.execute("SELECT SUM(CHAR_LENGTH(описание)) FROM entries WHERE описание IS NOT NULL")
        result = self.cursor.fetchone()[0]
        total_chars = result if result is not None else 0
        if total_chars > 5000:
            self.unlock_achievement("Словобог", user_id)

    def unlock_achievement(self, achievement_name, user_id):
        try:
            self.cursor.execute("SELECT id FROM achievements WHERE название = %s", (achievement_name,))
            achievement_id = self.cursor.fetchone()[0]

            self.cursor.execute("SELECT 1 FROM user_achievements WHERE user_id = %s AND achievement_id = %s",
                                (user_id, achievement_id))
            if not self.cursor.fetchone():
                self.cursor.execute("INSERT INTO user_achievements (user_id, achievement_id) VALUES (%s, %s)",
                                    (user_id, achievement_id))
                self.conn.commit()
                messagebox.showinfo("Достижение", f"Получено достижение: {achievement_name}")
        except Exception as e:
            print(f"Ошибка разблокировки достижения: {e}")

    def create_research_map_tab(self):
        map_tab = ttk.Frame(self.notebook)
        self.notebook.add(map_tab, text="Моя исследовательская карта")

        frame = ttk.LabelFrame(map_tab, text="Статистика")
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.research_text = tk.Text(frame, wrap=tk.WORD, width=80, height=25)
        self.research_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        ttk.Button(frame, text="Обновить статистику", command=self.update_research_map).pack(pady=10)

    def update_research_map(self):
        try:
            self.cursor.execute("""
                SELECT k.keyword, COUNT(ek.entry_id) as count
                FROM keywords k 
                LEFT JOIN entry_keywords ek ON k.id = ek.keyword_id 
                GROUP BY k.keyword 
                ORDER BY count DESC, k.keyword
            """)
            keywords_data = self.cursor.fetchall()

            self.cursor.execute("""
                SELECT соавторы FROM entries 
                WHERE соавторы IS NOT NULL AND соавторы != '' AND соавторы != ' '
            """)
            coauthors_data = self.cursor.fetchall()

            coauthor_count = {}
            for row in coauthors_data:
                if row[0]:
                    authors = [a.strip() for a in row[0].split(",") if a.strip()]
                    for author in authors:
                        coauthor_count[author] = coauthor_count.get(author, 0) + 1

            sorted_coauthors = sorted(coauthor_count.items(), key=lambda x: x[1], reverse=True)

            text = "Ключевые слова:\n"
            for keyword, count in keywords_data:
                if count > 0:
                    text += f"  {keyword} — {count} записи\n"

            if not any(count > 0 for _, count in keywords_data):
                text += "  (ключевые слова еще не добавлены)\n"

            text += "\nСоавторы:\n"
            for coauthor, count in sorted_coauthors:
                if coauthor:
                    text += f"  {coauthor} — {count} работы\n"

            if not sorted_coauthors:
                text += "  (соавторы еще не добавлены)\n"

            self.research_text.delete("1.0", tk.END)
            self.research_text.insert("1.0", text)

        except Exception as e:
            self.research_text.delete("1.0", tk.END)
            self.research_text.insert("1.0", f"Ошибка при загрузке данных: {str(e)}")

    def create_achievements_tab(self):
        achievements_tab = ttk.Frame(self.notebook)
        self.notebook.add(achievements_tab, text="Достижения")

        frame = ttk.LabelFrame(achievements_tab, text="Полученные достижения")
        frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.achievements_text = tk.Text(frame, wrap=tk.WORD, width=80, height=25)
        self.achievements_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        ttk.Button(frame, text="Обновить список", command=self.update_achievements).pack(pady=10)

    def update_achievements(self):
        try:
            self.cursor.execute("""
                SELECT a.название, a.описание, ua.получено 
                FROM achievements a 
                LEFT JOIN user_achievements ua ON a.id = ua.achievement_id AND ua.user_id = 1
                ORDER BY a.id
            """)
            achievements = self.cursor.fetchall()

            text = "Все достижения:\n\n"
            for name, desc, date in achievements:
                if date:
                    date_str = date.strftime("%Y-%m-%d %H:%M")
                    status = f"✓ Получено: {date_str}"
                else:
                    status = "✗ Еще не получено"

                text += f"• {name}\n  {desc}\n  {status}\n\n"

            self.achievements_text.delete("1.0", tk.END)
            self.achievements_text.insert("1.0", text)

        except Exception as e:
            self.achievements_text.delete("1.0", tk.END)
            self.achievements_text.insert("1.0", f"Ошибка при загрузке данных: {str(e)}")

    def create_competencies_tab(self):
        competencies_tab = ttk.Frame(self.notebook)
        self.notebook.add(competencies_tab, text="Компетенции")

        top_frame = ttk.Frame(competencies_tab)
        top_frame.pack(fill=tk.X, padx=10, pady=5)

        ttk.Label(top_frame, text="Специальность:").pack(side=tk.LEFT, padx=5)
        self.specialty_combo = ttk.Combobox(top_frame, values=list(self.specialties_data.keys()), state="readonly",
                                            width=30)
        self.specialty_combo.pack(side=tk.LEFT, padx=5)

        ttk.Button(top_frame, text="Загрузить компетенции", command=self.load_competencies).pack(side=tk.LEFT, padx=5)

        main_frame = ttk.Frame(competencies_tab)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=5)

        left_frame = ttk.LabelFrame(main_frame, text="Мои компетенции")
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=5)

        self.competencies_text = tk.Text(left_frame, wrap=tk.WORD, width=40, height=20)
        self.competencies_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        right_frame = ttk.LabelFrame(main_frame, text="Рекомендации")
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=5)

        self.recommendations_text = tk.Text(right_frame, wrap=tk.WORD, width=40, height=20)
        self.recommendations_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        ttk.Button(main_frame, text="Обновить профиль", command=self.update_competencies).pack(pady=10)

        if self.specialties_data:
            self.specialty_combo.set(list(self.specialties_data.keys())[0])
            self.load_competencies()

    def load_competencies(self):
        specialty = self.specialty_combo.get()
        if specialty and specialty in self.specialties_data:
            self.current_specialty = specialty
            self.competencies = []

            try:
                self.cursor.execute("DELETE FROM competencies")

                competencies_list = self.specialties_data[specialty].get("competencies", [])

                for comp in competencies_list:
                    self.cursor.execute(
                        "INSERT INTO competencies (название, категория) VALUES (%s, %s) RETURNING id",
                        (comp['название'], comp['категория'])
                    )
                    comp_id = self.cursor.fetchone()[0]
                    self.competencies.append(
                        {'id': comp_id, 'название': comp['название'], 'категория': comp['категория']})

                self.conn.commit()
                self.create_competencies_widgets()
                self.update_competencies()

                messagebox.showinfo("Успех", f"Загружены компетенции для специальности: {specialty}")

            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка загрузки компетенций: {str(e)}")

    def update_competencies(self):
        if not self.competencies:
            self.competencies_text.delete("1.0", tk.END)
            self.competencies_text.insert("1.0", "Сначала загрузите компетенции для специальности")
            self.recommendations_text.delete("1.0", tk.END)
            return

        try:
            self.cursor.execute("""
                SELECT c.id, c.название, c.категория, AVG(ec.уровень) as avg_level
                FROM competencies c
                LEFT JOIN entry_competencies ec ON c.id = ec.competency_id
                GROUP BY c.id, c.название, c.категория
                ORDER BY c.категория, c.название
            """)
            comp_data = self.cursor.fetchall()

            text = "Средний уровень по компетенциям:\n\n"
            weak_zones = []

            for comp_id, name, category, avg_level in comp_data:
                level = round(float(avg_level or 0), 1)
                text += f"• {name} ({category}): {level}/5.0\n"
                if level < 3 and level > 0:
                    weak_zones.append((name, level))

            if weak_zones:
                text += "\nСлабые зоны (уровень < 3):\n"
                for name, level in weak_zones:
                    text += f"• {name}: {level}/5.0\n"

            self.competencies_text.delete("1.0", tk.END)
            self.competencies_text.insert("1.0", text)

            self.generate_recommendations(comp_data)

        except Exception as e:
            self.competencies_text.delete("1.0", tk.END)
            self.competencies_text.insert("1.0", f"Ошибка при загрузке компетенций: {str(e)}")

    def generate_recommendations(self, comp_data):
        recommendations = []

        for comp_id, name, category, avg_level in comp_data:
            level = float(avg_level or 0)

            if level == 0:
                recommendations.append(
                    f"Вы еще не оценивали компетенцию '{name}'. Добавьте запись с этой компетенцией.")
            elif level < 2:
                if "Программирование" in name:
                    recommendations.append(
                        f"Компетенция '{name}' требует развития. Рекомендуем решать задачи на LeetCode или Codewars.")
                elif "БД" in name:
                    recommendations.append(
                        f"Компетенция '{name}' низкая. Рекомендуем пройти курс по SQL и NoSQL базам данных.")
                elif "Презентация" in name:
                    recommendations.append(
                        f"Компетенция '{name}' слабая. Выступите с докладом на студенческой конференции.")
                else:
                    recommendations.append(
                        f"Компетенция '{name}' требует серьезного развития. Рекомендуем пройти соответствующий курс.")
            elif level < 3:
                recommendations.append(
                    f"Уровень компетенции '{name}' ниже среднего. Рекомендуем практиковаться в этой области.")
            elif level >= 4.5:
                recommendations.append(f"Компетенция '{name}' развита отлично! Можете делиться опытом с другими.")

        if not recommendations:
            recommendations.append("Все компетенции развиты хорошо! Продолжайте в том же духе.")

        rec_text = "Персонализированные рекомендации:\n\n"
        for i, rec in enumerate(recommendations, 1):
            rec_text += f"{i}. {rec}\n"

        self.recommendations_text.delete("1.0", tk.END)
        self.recommendations_text.insert("1.0", rec_text)

    def create_goals_tab(self):
        goals_tab = ttk.Frame(self.notebook)
        self.notebook.add(goals_tab, text="Цели на семестр")

        left_frame = ttk.LabelFrame(goals_tab, text="Новая цель")
        left_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        ttk.Label(left_frame, text="Тип цели:").grid(row=0, column=0, sticky=tk.W, padx=5, pady=5)
        self.goal_type_combo = ttk.Combobox(left_frame, values=["Добавить записи", "Поднять компетенцию"],
                                            state="readonly", width=25)
        self.goal_type_combo.grid(row=0, column=1, padx=5, pady=5)
        self.goal_type_combo.set("Добавить записи")
        self.goal_type_combo.bind("<<ComboboxSelected>>", self.on_goal_type_change)

        ttk.Label(left_frame, text="Описание:").grid(row=1, column=0, sticky=tk.W, padx=5, pady=5)
        self.goal_desc_entry = ttk.Entry(left_frame, width=30)
        self.goal_desc_entry.grid(row=1, column=1, padx=5, pady=5)

        ttk.Label(left_frame, text="Целевое значение:").grid(row=2, column=0, sticky=tk.W, padx=5, pady=5)
        self.goal_value_entry = ttk.Entry(left_frame, width=30)
        self.goal_value_entry.grid(row=2, column=1, padx=5, pady=5)
        self.goal_value_entry.insert(0, "2")

        ttk.Label(left_frame, text="Компетенция:").grid(row=3, column=0, sticky=tk.W, padx=5, pady=5)
        self.goal_comp_combo = ttk.Combobox(left_frame, width=28, state="disabled")
        self.goal_comp_combo.grid(row=3, column=1, padx=5, pady=5)

        ttk.Button(left_frame, text="Добавить цель", command=self.add_goal).grid(row=4, column=0, columnspan=2, pady=20)

        right_frame = ttk.LabelFrame(goals_tab, text="Мои цели")
        right_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=10, pady=10)

        self.goals_text = tk.Text(right_frame, wrap=tk.WORD, width=50, height=20)
        self.goals_text.pack(fill=tk.BOTH, expand=True, padx=5, pady=5)

        ttk.Button(right_frame, text="Обновить прогресс", command=self.update_goals).pack(pady=10)
        ttk.Button(right_frame, text="Удалить все цели", command=self.delete_all_goals).pack(pady=5)

    def on_goal_type_change(self, event=None):
        if self.goal_type_combo.get() == "Поднять компетенцию":
            self.goal_comp_combo['state'] = 'normal'
            comp_names = [comp['название'] for comp in self.competencies]
            self.goal_comp_combo['values'] = comp_names
            if comp_names:
                self.goal_comp_combo.set(comp_names[0])
        else:
            self.goal_comp_combo['state'] = 'disabled'
            self.goal_comp_combo.set('')

    def add_goal(self):
        goal_type = self.goal_type_combo.get()
        description = self.goal_desc_entry.get()

        if not description:
            messagebox.showwarning("Предупреждение", "Введите описание цели")
            return

        try:
            target_value = int(self.goal_value_entry.get())
            if target_value <= 0:
                raise ValueError
        except:
            messagebox.showwarning("Предупреждение", "Введите положительное числовое значение")
            return

        if goal_type == "Поднять компетенцию":
            competency = self.goal_comp_combo.get()
            if not competency:
                messagebox.showwarning("Предупреждение", "Выберите компетенцию")
                return
            description = f"{description} ({competency})"

        try:
            self.cursor.execute(
                "INSERT INTO goals (описание, тип, цель_значение) VALUES (%s, %s, %s)",
                (description, goal_type, target_value)
            )
            self.conn.commit()

            self.goal_desc_entry.delete(0, tk.END)
            messagebox.showinfo("Успех", "Цель добавлена")

            self.update_goals()

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при добавлении цели: {str(e)}")

    def update_goals(self):
        try:
            self.cursor.execute("SELECT * FROM goals ORDER BY id")
            goals = self.cursor.fetchall()

            text = ""
            for goal in goals:
                id, description, goal_type, target, current, completed = goal

                if goal_type == "Добавить записи":
                    self.cursor.execute("SELECT COUNT(*) FROM entries")
                    current = self.cursor.fetchone()[0]
                elif goal_type == "Поднять компетенцию":
                    comp_name = description.split("(")[-1].rstrip(")") if "(" in description else ""
                    if comp_name:
                        self.cursor.execute("""
                            SELECT AVG(ec.уровень) 
                            FROM competencies c
                            JOIN entry_competencies ec ON c.id = ec.competency_id
                            WHERE c.название = %s
                        """, (comp_name,))
                        result = self.cursor.fetchone()[0]
                        current = round(float(result or 0), 1)
                    else:
                        current = 0

                self.cursor.execute("UPDATE goals SET текущее_значение = %s WHERE id = %s", (current, id))

                progress = f"{current} из {target}"
                if current >= target:
                    status = "✓ Выполнено"
                    self.cursor.execute("UPDATE goals SET завершено = TRUE WHERE id = %s", (id,))
                else:
                    status = "→ В процессе"
                    self.cursor.execute("UPDATE goals SET завершено = FALSE WHERE id = %s", (id,))

                text += f"Цель: {description}\nПрогресс: {progress} {status}\n\n"

            self.conn.commit()
            self.goals_text.delete("1.0", tk.END)
            self.goals_text.insert("1.0", text if text else "Цели еще не добавлены")

        except Exception as e:
            self.goals_text.delete("1.0", tk.END)
            self.goals_text.insert("1.0", f"Ошибка при обновлении целей: {str(e)}")

    def delete_all_goals(self):
        if messagebox.askyesno("Подтверждение", "Удалить все цели?"):
            try:
                self.cursor.execute("DELETE FROM goals")
                self.conn.commit()
                self.update_goals()
                messagebox.showinfo("Успех", "Все цели удалены")
            except Exception as e:
                messagebox.showerror("Ошибка", f"Ошибка при удалении: {str(e)}")

    def export_to_word(self):
        try:
            doc = Document()

            title = doc.add_heading('Отчет по портфолио', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_heading('1. Все записи', level=1)
            self.cursor.execute("SELECT * FROM entries ORDER BY дата DESC")
            entries = self.cursor.fetchall()

            if entries:
                for entry in entries:
                    id, title_text, entry_type, date, description, coauthors = entry
                    doc.add_heading(f"{title_text} ({entry_type})", level=2)
                    doc.add_paragraph(f"Дата: {date}")
                    if coauthors:
                        doc.add_paragraph(f"Соавторы: {coauthors}")
                    if description:
                        doc.add_paragraph(f"Описание: {description}")
                    doc.add_paragraph()
            else:
                doc.add_paragraph("Записей нет")

            doc.add_heading('2. Исследовательская карта', level=1)
            self.update_research_map()
            map_text = self.research_text.get("1.0", tk.END).strip()
            for line in map_text.split('\n'):
                if line:
                    doc.add_paragraph(line)

            doc.add_heading('3. Профиль компетенций', level=1)
            self.update_competencies()
            comp_text = self.competencies_text.get("1.0", tk.END).strip()
            for line in comp_text.split('\n'):
                if line:
                    doc.add_paragraph(line)

            doc.add_heading('4. Рекомендации', level=1)
            self.generate_recommendations([])
            rec_text = self.recommendations_text.get("1.0", tk.END).strip()
            for line in rec_text.split('\n'):
                if line:
                    doc.add_paragraph(line)

            doc.add_heading('5. Достижения', level=1)
            self.update_achievements()
            ach_text = self.achievements_text.get("1.0", tk.END).strip()
            for line in ach_text.split('\n'):
                if line:
                    doc.add_paragraph(line)

            doc.add_heading('6. Цели на семестр', level=1)
            self.update_goals()
            goals_text = self.goals_text.get("1.0", tk.END).strip()
            for line in goals_text.split('\n'):
                if line:
                    doc.add_paragraph(line)

            filename = f"portfolio_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(filename)

            messagebox.showinfo("Успех", f"Отчет сохранен в файл: {filename}")

        except Exception as e:
            messagebox.showerror("Ошибка", f"Ошибка при экспорте: {str(e)}")


if __name__ == "__main__":
    root = tk.Tk()
    app = PortfolioApp(root)
    root.mainloop()

